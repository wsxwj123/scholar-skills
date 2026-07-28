#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
博士论文质量自检工具

功能：
1. 检查字数是否达标
2. 检查格式是否规范
3. 检查图表编号是否连续
4. 检查是否有列表项（应为段落）
5. 生成质量报告

作者：Sci2Doc Team
日期：2024-03-15
"""

import sys as _sys
try:  # Windows GBK 控制台/管道捕获下 emoji print 防 UnicodeEncodeError
    _sys.stdout.reconfigure(encoding="utf-8")
    _sys.stderr.reconfigure(encoding="utf-8")
except Exception:
    pass
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import argparse
import sys
import os
import re
import json
from datetime import datetime

try:
    from thesis_profile import build_format_render_context, load_profile
except Exception:  # pragma: no cover
    script_dir = os.path.dirname(os.path.abspath(__file__))
    if script_dir not in sys.path:
        sys.path.insert(0, script_dir)
    from thesis_profile import build_format_render_context, load_profile

try:
    from shared_utils import normalize_text, heading_level, classify_heading, infer_project_root_for_profile, classify_paragraphs_with_chapter_scope
except ImportError:  # pragma: no cover
    script_dir = os.path.dirname(os.path.abspath(__file__))
    if script_dir not in sys.path:
        sys.path.insert(0, script_dir)
    from shared_utils import normalize_text, heading_level, classify_heading, infer_project_root_for_profile, classify_paragraphs_with_chapter_scope

try:
    from abbreviation_registry import load_registry, extract_abbreviations, validate_cross_references
except Exception:  # pragma: no cover
    _script_dir = os.path.dirname(os.path.abspath(__file__))
    if _script_dir not in sys.path:
        sys.path.insert(0, _script_dir)
    try:
        from abbreviation_registry import load_registry, extract_abbreviations, validate_cross_references
    except ImportError:
        load_registry = None
        extract_abbreviations = None
        validate_cross_references = None

try:
    from reference_renderer import validate_all as _validate_references
except Exception:  # pragma: no cover
    _ref_script_dir = os.path.dirname(os.path.abspath(__file__))
    if _ref_script_dir not in sys.path:
        sys.path.insert(0, _ref_script_dir)
    try:
        from reference_renderer import validate_all as _validate_references
    except ImportError:
        _validate_references = None


def _pending_template_payload(docx_path):
    try:
        project_root = infer_project_root_for_profile(docx_path)
        profile, _ = load_profile(project_root)
    except Exception:
        return None
    format_profile = profile.get("format_profile", {}) if isinstance(profile, dict) else {}
    if str(format_profile.get("status", "")).strip() != "pending_template":
        return None
    return {
        "success": False,
        "error": "pending_template",
        "file_path": docx_path,
        "project_root": project_root,
        "format_profile": format_profile,
        "message": "当前项目处于 pending_template，禁止执行格式验收。请先补齐自定义院校格式模板要求。",
    }


def _load_render_context_for_docx(docx_path):
    fallback = build_format_render_context({})
    try:
        project_root = infer_project_root_for_profile(docx_path)
        profile, _ = load_profile(project_root)
    except Exception:
        return fallback
    format_profile = profile.get("format_profile", {}) if isinstance(profile, dict) else {}
    return build_format_render_context(format_profile)


def line_spacing_pt(paragraph):
    value = paragraph.paragraph_format.line_spacing
    if value is None:
        return None
    if hasattr(value, "pt"):
        return float(value.pt)
    if isinstance(value, (int, float)):
        if value > 1000:
            return float(value) / 12700.0
        return float(value)
    return None


def check_word_count(doc, body_target_chars=80000, review_target_chars=0, review_in_scope=False, extra_exclude=None):
    """检查字数是否达标（使用章作用域继承分类，摘要计入正文）"""
    issues = []

    total_chinese = 0
    review_chinese = 0

    para_stream = (
        (para.text.strip(), heading_level(getattr(para.style, "name", "")))
        for para in doc.paragraphs
        if para.text.strip()
    )

    for text, lvl, sec_type in classify_paragraphs_with_chapter_scope(para_stream, extra_exclude=extra_exclude):
        if lvl is not None:
            continue  # 标题行本身不计字数

        chars = sum(1 for char in text if '\u4e00' <= char <= '\u9fff')
        if sec_type == "review":
            # \u7efc\u8ff0/\u7eea\u8bba\u7ae0\u8ba1\u5165\u6b63\u6587\u5b57\u6570\uff0c\u628a\u51d1\u5b57\u6570\u538b\u529b\u4ece\u7814\u7a76\u7ae0\u5206\u644a\u51fa\u53bb\uff0c\u7f13\u89e3\u903c\u9020\u6570\u636e\u3002
            # review_chinese \u4ecd\u5355\u72ec\u7edf\u8ba1\u4f9b"\u7efc\u8ff0\u5b57\u6570"\u5c55\u793a\u4e0e review_target \u63d0\u793a\u3002
            review_chinese += chars
            total_chinese += chars
        elif sec_type in {"references", "toc", "acknowledgement", "appendix", "achievements", "declaration", "abbreviation_table"}:
            continue
        else:
            # body and abstract both count
            total_chinese += chars
    
    # 检查正文字数
    body_target_chars = max(1, int(body_target_chars or 80000))
    review_target_chars = max(0, int(review_target_chars or 0))

    if total_chinese < body_target_chars:
        issues.append({
            # 软提示（非硬门）：字数地板是目标不是硬拦，避免逼 AT 靠编数据/灌水凑字。
            'level': 'warning',
            'category': '字数',
            'message': f'正文字数不足：{total_chinese} / {body_target_chars:,} 字（提示，非硬门）',
            'suggestion': f'用真实材料扩展至 {body_target_chars:,} 字目标；宁可少写也不得编造数据凑字'
        })
    
    # 检查综述字数
    if review_in_scope and review_target_chars > 0 and review_chinese > 0 and review_chinese < review_target_chars:
        issues.append({
            'level': 'warning',
            'category': '字数',
            'message': f'综述字数不足：{review_chinese} / {review_target_chars:,} 字',
            'suggestion': f'建议扩展综述部分至 {review_target_chars:,} 字以上'
        })
    
    return issues, {
        'body_words': total_chinese,
        'review_words': review_chinese,
        'body_target_chars': body_target_chars,
        'review_target_chars': review_target_chars,
        'review_in_scope': bool(review_in_scope),
    }


def check_heading_levels(doc):
    """检查标题层级是否规范"""
    issues = []
    prev_level = 0
    
    for i, para in enumerate(doc.paragraphs):
        level = heading_level(getattr(para.style, "name", ""))
        if level is not None:
            
            # 检查是否超过三级
            if level > 3:
                issues.append({
                    'level': 'error',
                    'category': '标题层级',
                    'location': f'第 {i+1} 段',
                    'message': f'标题层级过深：{level} 级（{para.text[:30]}...）',
                    'suggestion': '建议标题层级最多三级'
                })
            
            # 检查是否跳级（如 1 → 3）
            if level - prev_level > 1:
                issues.append({
                    'level': 'warning',
                    'category': '标题层级',
                    'location': f'第 {i+1} 段',
                    'message': f'标题层级跳跃：从 {prev_level} 级跳到 {level} 级',
                    'suggestion': '建议按顺序设置标题层级'
                })
            
            prev_level = level
    
    return issues


def check_figure_numbering(doc):
    """检查图表编号是否连续规范"""
    issues = []

    # 匹配图编号：图 1-1, 图 2-3 等
    figure_pattern = re.compile(r'图\s*(\d+)-(\d+)')
    table_pattern = re.compile(r'表\s*(\d+)-(\d+)')
    # A1 公式编号：公式(2-3) / 式(2-3) / 公式 2-3 / 式 2-3（按章 chapter-number）
    formula_pattern = re.compile(r'(?:公式|式)\s*[（(]?\s*(\d+)\s*-\s*(\d+)\s*[）)]?')
    
    figures = []
    tables = []
    formulas = []

    for i, para in enumerate(doc.paragraphs):
        text = para.text

        # 查找图编号
        for match in figure_pattern.finditer(text):
            chapter = int(match.group(1))
            number = int(match.group(2))
            figures.append({
                'chapter': chapter,
                'number': number,
                'location': i + 1,
                'text': text[:50]
            })

        # 查找表编号
        for match in table_pattern.finditer(text):
            chapter = int(match.group(1))
            number = int(match.group(2))
            tables.append({
                'chapter': chapter,
                'number': number,
                'location': i + 1,
                'text': text[:50]
            })

        # 查找公式编号（A1）
        for match in formula_pattern.finditer(text):
            chapter = int(match.group(1))
            number = int(match.group(2))
            formulas.append({
                'chapter': chapter,
                'number': number,
                'location': i + 1,
                'text': text[:50]
            })
    
    # 检查图编号连续性
    figures_by_chapter = {}
    for fig in figures:
        chapter = fig['chapter']
        if chapter not in figures_by_chapter:
            figures_by_chapter[chapter] = []
        figures_by_chapter[chapter].append(fig)
    
    for chapter, figs in figures_by_chapter.items():
        # 按 (chapter, number) 去重：同一图号被题注+正文多次引用只算一次，
        # 否则排序后 expected=i+1 必错位，产生假"不连续"warning。
        seen_numbers = set()
        figs_unique = []
        for fig in figs:
            if fig['number'] not in seen_numbers:
                seen_numbers.add(fig['number'])
                figs_unique.append(fig)
        figs_sorted = sorted(figs_unique, key=lambda x: x['number'])
        for i, fig in enumerate(figs_sorted):
            expected = i + 1
            if fig['number'] != expected:
                issues.append({
                    'level': 'warning',
                    'category': '图表编号',
                    'location': f'第 {fig["location"]} 段',
                    'message': f'第 {chapter} 章图编号不连续：图 {chapter}-{fig["number"]}',
                    'suggestion': f'应为 图 {chapter}-{expected}'
                })
    
    # 检查表编号连续性（同理）
    tables_by_chapter = {}
    for tab in tables:
        chapter = tab['chapter']
        if chapter not in tables_by_chapter:
            tables_by_chapter[chapter] = []
        tables_by_chapter[chapter].append(tab)
    
    for chapter, tabs in tables_by_chapter.items():
        # 按 (chapter, number) 去重：同一表号被题注+正文多次引用只算一次，
        # 否则排序后 expected=i+1 必错位，产生假"不连续"warning。
        seen_numbers = set()
        tabs_unique = []
        for tab in tabs:
            if tab['number'] not in seen_numbers:
                seen_numbers.add(tab['number'])
                tabs_unique.append(tab)
        tabs_sorted = sorted(tabs_unique, key=lambda x: x['number'])
        for i, tab in enumerate(tabs_sorted):
            expected = i + 1
            if tab['number'] != expected:
                issues.append({
                    'level': 'warning',
                    'category': '图表编号',
                    'location': f'第 {tab["location"]} 段',
                    'message': f'第 {chapter} 章表编号不连续：表 {chapter}-{tab["number"]}',
                    'suggestion': f'应为 表 {chapter}-{expected}'
                })

    # 检查公式编号连续性（A1，同表/图去重写法）
    formulas_by_chapter = {}
    for fml in formulas:
        chapter = fml['chapter']
        if chapter not in formulas_by_chapter:
            formulas_by_chapter[chapter] = []
        formulas_by_chapter[chapter].append(fml)

    for chapter, fmls in formulas_by_chapter.items():
        # 按 (chapter, number) 去重：同一公式号被定义+正文多次引用只算一次，
        # 否则排序后 expected=i+1 必错位，产生假"不连续"warning。
        seen_numbers = set()
        fmls_unique = []
        for fml in fmls:
            if fml['number'] not in seen_numbers:
                seen_numbers.add(fml['number'])
                fmls_unique.append(fml)
        fmls_sorted = sorted(fmls_unique, key=lambda x: x['number'])
        for i, fml in enumerate(fmls_sorted):
            expected = i + 1
            if fml['number'] != expected:
                issues.append({
                    'level': 'warning',
                    'category': '公式编号',
                    'location': f'第 {fml["location"]} 段',
                    'message': f'第 {chapter} 章公式编号不连续：公式 {chapter}-{fml["number"]}',
                    'suggestion': f'应为 公式 {chapter}-{expected}'
                })

    return issues


def check_figure_map_consistency(doc, project_root=None):
    """
    交叉验证 figure_map.json 与 docx 中的图编号。

    检查：
    1. docx 中引用的图编号是否都在 figure_map 中注册
    2. figure_map 中注册的图编号是否都在 docx 中被引用
    """
    issues = []
    if project_root is None:
        return issues

    figure_map_path = os.path.join(os.path.abspath(project_root), 'figure_map.json')
    if not os.path.exists(figure_map_path):
        return issues

    try:
        with open(figure_map_path, 'r', encoding='utf-8') as f:
            figure_map = json.load(f)
        if not isinstance(figure_map, dict):
            return issues
    except (json.JSONDecodeError, OSError):
        return issues

    registered_ids = set(figure_map.keys())

    # 收集 docx 中所有图编号
    figure_pattern = re.compile(r'图\s*(\d+)\s*-\s*(\d+)')
    doc_figure_ids = set()
    for para in doc.paragraphs:
        for m in figure_pattern.finditer(para.text):
            cn_id = f"图{m.group(1)}-{m.group(2)}"
            doc_figure_ids.add(cn_id)

    unregistered = sorted(doc_figure_ids - registered_ids)
    unreferenced = sorted(registered_ids - doc_figure_ids)

    for cn_id in unregistered:
        issues.append({
            'level': 'warning',
            'category': '图编号映射',
            'message': f'文档中引用了 {cn_id}，但未在 figure_map.json 中注册',
            'suggestion': f'运行 figure_registry.py register 注册该图映射'
        })

    for cn_id in unreferenced:
        issues.append({
            'level': 'info',
            'category': '图编号映射',
            'message': f'{cn_id} 已注册但未在文档中引用',
            'suggestion': '确认该图是否已从文档中移除'
        })

    return issues


def check_bullet_points(doc):
    """检查是否有列表项（应全部为段落）"""
    issues = []
    
    for i, para in enumerate(doc.paragraphs):
        # 检查是否为列表样式
        if 'List' in para.style.name:
            issues.append({
                'level': 'error',
                'category': '格式',
                'location': f'第 {i+1} 段',
                'message': f'发现列表项：{para.text[:30]}...',
                'suggestion': '论文要求段落式写作，不允许使用列表项'
            })
        
        # 检查是否有列表标记（简单检测）
        text = para.text.strip()
        if text and (
            text.startswith('• ') or
            text.startswith('- ') or
            text.startswith('* ') or
            re.match(r'^\d+\.\s', text) or
            re.match(r'^\(\d+\)\s', text)
        ):
            issues.append({
                'level': 'warning',
                'category': '格式',
                'location': f'第 {i+1} 段',
                'message': f'疑似列表项：{text[:30]}...',
                'suggestion': '请确认是否为列表项，如是请改为段落形式'
            })
    
    return issues


def check_reference_count(doc, min_reference_count=80):
    """检查参考文献数量"""
    issues = []
    
    # 查找参考文献章节
    in_references = False
    ref_count = 0
    
    for para in doc.paragraphs:
        text = para.text.strip()

        if not text:
            continue

        lvl = heading_level(getattr(para.style, "name", ""))
        if lvl is not None:
            if classify_heading(text) == "references":
                in_references = True
            elif in_references:
                # 下一标题即离开参考文献章节
                in_references = False
            continue

        if in_references:
            # 匹配 [1], [2] 等编号
            if re.match(r'^\[\d+\]', text) or re.match(r'^\d+\.\s', text):
                ref_count += 1
    
    min_reference_count = int(min_reference_count if min_reference_count is not None else 80)
    if min_reference_count <= 0:
        return issues, ref_count
    if ref_count < min_reference_count:
        issues.append({
            'level': 'error',
            'category': '参考文献',
            'message': f'参考文献数量不足：{ref_count} / {min_reference_count} 篇',
            'suggestion': f'当前配置要求参考文献不少于 {min_reference_count} 篇'
        })
    
    return issues, ref_count


def is_chapter_heading_text(text):
    t = (text or "").strip()
    return bool(re.match(r"^第[一二三四五六七八九十百千万零〇0-9]+章", t))


_CITE_MARKER_RE = re.compile(r"\[\d+(?:\s*[-–,，]\s*\d+)*\]")


def check_chapter_reference_distribution(doc, per_chapter_ref_floor):
    """按章软文献统计（warning，不阻断）：绪论/文献综述章 [n] 引用数 < intro_review 地板、
    研究/实验章 < research 地板 → warning。阈值来自 thesis_profile 的 per_chapter_ref_floor
    （硕/博分档）。总量硬门仍由 check_reference_count 负责。结论章不设文献地板。"""
    issues = []
    if not isinstance(per_chapter_ref_floor, dict):
        return issues
    intro_floor = int(per_chapter_ref_floor.get("intro_review", 0) or 0)
    research_floor = int(per_chapter_ref_floor.get("research", 0) or 0)
    if intro_floor <= 0 and research_floor <= 0:
        return issues

    skip_kinds = {"references", "toc", "abstract", "acknowledgement", "appendix",
                  "achievements", "declaration", "abbreviation_table"}
    chapters = []  # {title, bucket, count}
    cur = None

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        lvl = heading_level(getattr(para.style, "name", ""))
        if lvl is not None and is_chapter_heading_text(text):
            if cur is not None:
                chapters.append(cur)
            kind = classify_heading(text)
            norm = normalize_text(text)
            if kind in skip_kinds:
                cur = None
                continue
            if kind == "review" or any(k in norm for k in ("绪论", "引言", "文献综述", "综述", "研究背景", "研究进展")):
                bucket = "intro_review"
            elif any(k in norm for k in ("总结", "结论", "展望", "结束语")):
                bucket = None  # 结论章不设文献地板
            else:
                bucket = "research"
            cur = {"title": text, "bucket": bucket, "count": 0} if bucket else None
            continue
        if cur is not None and lvl is None:
            cur["count"] += len(_CITE_MARKER_RE.findall(text))
    if cur is not None:
        chapters.append(cur)

    for ch in chapters:
        floor = intro_floor if ch["bucket"] == "intro_review" else research_floor
        if floor <= 0:
            continue
        if ch["count"] < floor:
            label = "绪论/文献综述章" if ch["bucket"] == "intro_review" else "研究/实验章"
            issues.append({
                'level': 'warning',
                'category': '文献引用分布',
                'message': f'{ch["title"][:20]}（{label}）文献引用偏少：{ch["count"]} / 建议≥{floor} 处',
                'suggestion': f'{label}建议引用不少于 {floor} 处文献，当前仅 {ch["count"]} 处，请补充支撑或确认合理',
            })
    return issues


def check_reference_position(doc):
    """
    校验参考文献位置：
    1) 全文应只有一个参考文献标题
    2) 参考文献后不能再出现“第X章”正文章节标题
    """
    issues = []
    heading_items = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        lvl = heading_level(getattr(para.style, "name", ""))
        if lvl is None:
            continue
        heading_items.append((i, text, classify_heading(text)))

    ref_headings = [(idx, text) for idx, text, kind in heading_items if kind == "references"]
    if len(ref_headings) > 1:
        issues.append({
            'level': 'error',
            'category': '参考文献位置',
            'message': f'检测到多个参考文献标题：{len(ref_headings)} 处',
            'suggestion': '参考文献应全书统一放置，且仅保留一个参考文献章节'
        })
        return issues

    if len(ref_headings) == 1:
        ref_idx, _ = ref_headings[0]
        for idx, text, _kind in heading_items:
            if idx <= ref_idx:
                continue
            if is_chapter_heading_text(text):
                issues.append({
                    'level': 'error',
                    'category': '参考文献位置',
                    'location': f'第 {idx+1} 段',
                    'message': f'参考文献后出现章节标题：{text}',
                    'suggestion': '参考文献应位于正文章节之后，不能再进入新的“第X章”章节'
                })
                break

    return issues


# ---------------------------------------------------------------------------
# 正文内联引用格式检测
# ---------------------------------------------------------------------------

# 合法引用格式：[6] [6,7] [6-8] [6,7,9] [6-8,10] 等
_VALID_CITATION_RE = re.compile(
    r'\[(\d+(?:\s*[-–]\s*\d+)?(?:\s*,\s*\d+(?:\s*[-–]\s*\d+)?)*)\]'
)

# 宽松匹配：任何 [ 数字... ] 形态（用于发现格式错误的引用）
_LOOSE_CITATION_RE = re.compile(
    r'\[[\d,\-–\s]+\]'
)

# 不合法的常见错误模式
_BAD_CITATION_PATTERNS = [
    # 缺少逗号：[6 7]
    (re.compile(r'\[\d+\s+\d+\]'), '引用编号之间缺少逗号，应为 [X,Y] 格式'),
    # 中文逗号：[6，7]
    (re.compile(r'\[\d+\s*，\s*\d+'), '引用中使用了中文逗号，应使用英文逗号'),
    # 中文括号：（6）或【6】
    (re.compile(r'[（【]\d+(?:\s*[,，\-–]\s*\d+)*[）】]'), '引用使用了中文括号，应使用英文方括号 [X]'),
    # 上标格式残留：^[6] 或 <sup>[6]</sup>
    (re.compile(r'\^\[\d+'), '引用不应使用 markdown 上标语法'),
]


def check_inline_citations(doc):
    """
    检测正文中内联引用格式是否规范。

    合法格式：[6] [6,7] [6-8] [6,7,9-11]
    检测问题：
    - 格式错误的引用（中文逗号、缺逗号、中文括号等）
    - 引用编号不连续或逆序（如 [8,3]）
    - 引用编号超出参考文献范围（需配合 ref_count 使用）
    """
    issues = []
    in_references = False
    in_abstract = False
    para_idx = 0

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        lvl = heading_level(getattr(para.style, "name", ""))
        if lvl is not None:
            cls = classify_heading(text)
            if cls == "references":
                in_references = True
            elif in_references:
                in_references = False
            in_abstract = "摘要" in text or "abstract" in text.lower()
            continue

        # 跳过参考文献区域和摘要
        if in_references or in_abstract:
            continue

        para_idx += 1

        # 1) 检测错误格式
        for bad_re, msg in _BAD_CITATION_PATTERNS:
            for m in bad_re.finditer(text):
                snippet = text[max(0, m.start() - 10):m.end() + 10]
                issues.append({
                    'level': 'error',
                    'category': '引用格式',
                    'message': f'{msg}：...{snippet}...',
                    'suggestion': '正文引用应使用英文方括号+英文逗号，如 [1,2] [3-5]'
                })

        # 2) 检测合法引用中的编号问题
        for m in _VALID_CITATION_RE.finditer(text):
            inner = m.group(1)
            # 解析所有编号
            nums = []
            for part in re.split(r'\s*,\s*', inner):
                range_match = re.match(r'(\d+)\s*[-–]\s*(\d+)', part)
                if range_match:
                    start, end = int(range_match.group(1)), int(range_match.group(2))
                    if start >= end:
                        issues.append({
                            'level': 'error',
                            'category': '引用格式',
                            'message': f'引用范围逆序：[{inner}]，起始编号应小于结束编号',
                            'suggestion': f'应为 [{end}-{start}] 或拆分为独立编号'
                        })
                    nums.extend(range(start, end + 1))
                else:
                    nums.append(int(part.strip()))

            # 检查是否非递增（允许相等，不允许逆序）
            if len(nums) >= 2:
                for i in range(1, len(nums)):
                    if nums[i] < nums[i - 1]:
                        issues.append({
                            'level': 'warning',
                            'category': '引用格式',
                            'message': f'引用编号未按升序排列：[{inner}]',
                            'suggestion': '多引用应按编号升序排列，如 [3,5,7] 而非 [7,3,5]'
                        })
                        break

    return issues


# ---------------------------------------------------------------------------
# 写作风格检测
# ---------------------------------------------------------------------------

_STYLE_CHECKS = [
    # (compiled_regex, level, category, message, suggestion)
    (
        re.compile(r'——'),
        'error', '标点规范',
        '使用了破折号（——）',
        '用逗号、句号或重组句子代替破折号'
    ),
    (
        # 正文中的问号（排除引用他人原话的情况）
        re.compile(r'[？?]'),
        'warning', '陈述规范',
        '正文中出现问句',
        '正文应全部使用陈述句，不使用疑问句或反问句'
    ),
    (
        # 比喻词
        re.compile(r'如同|好比|仿佛|犹如|像[^素片].*?一样|恰似|宛如|宛若|好像(?!素)'),
        'error', '修辞规范',
        '使用了比喻修辞',
        '删除比喻表达，直接陈述事实'
    ),
    (
        # 常见比喻名词（...的桥梁/基石/钥匙/引擎/灯塔）
        re.compile(r'的(?:桥梁|基石|钥匙|引擎|灯塔|摇篮|沃土|温床|催化剂|助推器|风向标)'),
        'error', '修辞规范',
        '使用了比喻性名词',
        '用准确的功能描述替代比喻性名词'
    ),
    (
        # 主观夸大形容词
        re.compile(r'令人(?:惊讶|震惊|瞩目|振奋|鼓舞)|远超预期|出人意料|前所未有|史无前例|无与伦比|举世瞩目'),
        'warning', '客观性',
        '使用了主观色彩过强的表述',
        '用客观数据和事实描述结果，让读者自行判断'
    ),
    (
        # 过度书面化/生僻词
        re.compile(r'鉴于此|有鉴于此|兹|窃以为|殊为|诚然|毋庸置疑|不言而喻|众所周知'),
        'warning', '语言通俗性',
        '使用了过度书面化或套话表述',
        '用平实语言替代，如"因此""可以确认"等'
    ),
    (
        # scare quotes（恐惧引号）：双引号/引号包裹 2-10 字中文短语，暗示"新概念"
        # 保守检测：只报引号内全部是中文字符的情况（术语定义、固化术语通常有英文或括号）
        re.compile(r'"[一-鿿]{2,10}"|"[一-鿿]{2,10}"|“[一-鿿]{2,10}”'),
        'error', '引号规范',
        '疑似 scare quotes（引号包裹普通中文短语以暗示特殊概念）',
        '仅在术语首次定义、原文引用、已固化术语时使用引号；普通叙述不应使用引号强调'
    ),
    (
        # 解释性冒号（装饰句式）：纯中文词后跟中文冒号+中文内容（非图表标签/数值场景）
        # 排除：图N-N: / 表N-N: 前缀（前面是数字或图/表字符），或冒号后跟数字
        # 注意：保守检测，列表引导（"以下三点："）也会命中，属可接受误报
        re.compile(r'(?:[^表图\d]|^)([\u4e00-\u9fff]{2,15}：)(?=[\u4e00-\u9fff])'),
        'error', '标点规范',
        '疑似解释性冒号（"概念：解释"装饰句式）',
        '合法冒号场景：图表标签(表2-1：)、比例(1:2)、数值；正文装饰句式"概念：解释"应改写为陈述句'
    ),
]

# 排比检测：连续3个以上句子以相同模式开头
_PARALLELISM_MIN_COUNT = 3

# ---------------------------------------------------------------------------
# 中文 AI 禁词表（博论正文高频 AI 套话，参考 nsfc-proposal/scripts/humanizer_zh.py）
# 格式：(compiled_regex, category, message, suggestion)
# ---------------------------------------------------------------------------
_CN_AI_BANNED = [
    # 空洞修饰词
    (re.compile(r'至关重要|举足轻重|不可或缺'), '去AI-禁词', '空洞修饰词：至关重要/举足轻重/不可或缺', '用具体数据替代形容词'),
    (re.compile(r'具有重要的.{0,15}意义和.{0,15}价值'), '去AI-禁词', '套话：具有重要的...意义和...价值', '删除或改写为具体贡献陈述'),
    # 新闻体套话
    (re.compile(r'日益增长|蓬勃发展|方兴未艾'), '去AI-禁词', '新闻体套话：日益增长/蓬勃发展/方兴未艾', '替换为具体数据或趋势描述'),
    # 空洞动词
    (re.compile(r'深入探讨|系统研究|全面分析'), '去AI-禁词', '空洞动词：深入探讨/系统研究/全面分析', '改为具体研究行为，如"比较X与Y的差异"'),
    # 机械过渡
    (re.compile(r'综上所述|总而言之'), '去AI-禁词', '机械过渡：综上所述/总而言之', '用事实句结束段落，不用总结套话'),
    (re.compile(r'在此基础上|鉴于此'), '去AI-禁词', 'AI 过渡句：在此基础上/鉴于此', '直接写因果关系'),
    (re.compile(r'值得注意的是|需要指出的是'), '去AI-禁词', 'AI 提示语：值得注意的是/需要指出的是', '删除提示语，直接给结论/证据'),
    # 夸张修辞
    (re.compile(r'革命性的|颠覆性的|突破性的'), '去AI-禁词', '夸张修辞：革命性的/颠覆性的/突破性的', '用数据/事实说明程度'),
    # AI 句式：不是…而是
    (re.compile(r'不是.{1,25}?而是'), '去AI-禁词', 'AI 对比模板：不是…而是', '改为直接陈述句'),
    # AI 句式：不仅…而且
    (re.compile(r'不仅.{1,25}?[，,].{1,25}?而且'), '去AI-禁词', 'AI 递进模板：不仅…而且', '拆成两句事实陈述'),
    # 模糊表述
    (re.compile(r'大量研究表明'), '去AI-禁词', '模糊归因：大量研究表明', '改为明确作者+文献编号'),
    (re.compile(r'越来越多的证据'), '去AI-禁词', '模糊归因：越来越多的证据', '具体引用几项关键证据'),
    (re.compile(r'广泛应用'), '去AI-禁词', '模糊表述：广泛应用', '列出具体应用场景'),
]

# ---------------------------------------------------------------------------
# 中文单句字符数上限
# ---------------------------------------------------------------------------
_CN_SENT_MAX_CHARS = 50  # 单句 >50 中文字符报 cn_sentence_too_long
_CN_SENT_SHORT_MAX = 15  # ≤15 字视为"短句"
_CN_SENT_LONG_MIN = 30   # ≥30 字视为"长句"
_CN_SENT_MONOTONE_DIFF = 5  # 连续3句长度差 <5 字报单调警告

# ---------------------------------------------------------------------------
# 翻译腔软检测（SCI 英译中高频，全部 info 级软提示，不阻断）
# ---------------------------------------------------------------------------
_TE_NOMINALIZE = re.compile(r'进行[了]?|加以|予以|使得|得以')  # 名词化/虚化动词
_TE_BEI_PER_PARA = 3  # 单段"被"字达此数报被动句滥用
# 长定语链：短片段内连续 4 个及以上"的"（…的…的…的…的），典型英式前置定语堆叠
_TE_LONG_ATTR_CHAIN = re.compile(r'(?:[^的。！？；，]{1,8}的){4,}')


# ---------------------------------------------------------------------------
# 字符级检查（移植自 general-sci-writing/scripts/proofread.py，全部 WARN 级）
# D1 中文句内夹半角标点 / D2 应上下标却裸写 / F1 中文高置信错别字
# 强度：全部 warning，不阻断（与引用门禁 J5/J7 同级）
# ---------------------------------------------------------------------------

# D1：中文句内夹半角标点（高误报区，极度保守）。
# 只标"半角标点两侧紧邻汉字"这种高置信中文句内夹半角的情形。
# DOI/URL/数字/单位区间不会出现"汉字+半角+汉字"，故天然规避。
_HALFWIDTH_IN_CN = [
    (re.compile(r'([一-鿿]),([一-鿿])'), ',', '，', '中文句内半角逗号'),
    (re.compile(r'([一-鿿]);([一-鿿])'), ';', '；', '中文句内半角分号'),
    (re.compile(r'([一-鿿]):([一-鿿])'), ':', '：', '中文句内半角冒号'),
    (re.compile(r'([一-鿿])\(([一-鿿])'), '(', '（', '中文句内半角左括号'),
    (re.compile(r'([一-鿿])\)([一-鿿])'), ')', '）', '中文句内半角右括号'),
]

# D2：应上下标却裸写的化学式/标记。仅当未被 markdown 上下标(^..^/~..~)或
# HTML(<sup>/<sub>)包裹时才报；词边界 + 上下文约束尽量降低误报。
# 边界用 (?<![A-Za-z0-9])/(?![A-Za-z0-9]) 而非 \b：中文在 Unicode re 下属 \w，
# \b 在"汉字H2O汉字"处不成立，故用显式 ASCII 边界，确保中文环境也能抓到，
# 同时不误吞 XH2OY（前后接字母数字时不报）。
_SS = r'(?<![A-Za-z0-9])'   # 左：前面不是字母数字
_SE = r'(?![A-Za-z0-9])'    # 右：后面不是字母数字
_SUBSUP_PATTERNS = [
    (re.compile(_SS + r'H2O2' + _SE), 'H2O2 → H~2~O~2~（下标）'),
    (re.compile(_SS + r'H2O' + _SE), 'H2O → H~2~O（下标 2）'),
    (re.compile(_SS + r'CO2' + _SE), 'CO2 → CO~2~（下标 2）'),
    (re.compile(_SS + r'O2' + _SE), 'O2 → O~2~（下标 2）'),
    (re.compile(_SS + r'N2' + _SE), 'N2 → N~2~（下标 2）'),
    (re.compile(_SS + r'NH3' + _SE), 'NH3 → NH~3~（下标 3）'),
    (re.compile(_SS + r'SO2' + _SE), 'SO2 → SO~2~（下标 2）'),
    (re.compile(_SS + r'IC50' + _SE), 'IC50 → IC~50~（下标 50）'),
    (re.compile(_SS + r'EC50' + _SE), 'EC50 → EC~50~（下标 50）'),
    (re.compile(_SS + r'LD50' + _SE), 'LD50 → LD~50~（下标 50）'),
    (re.compile(_SS + r'(\d+(?:\.\d+)?)\s?(cm|mm|nm|μm|um|m)2' + _SE), '面积单位：如 cm2 → cm^2^（上标指数）'),
    (re.compile(_SS + r'(\d+(?:\.\d+)?)\s?(cm|mm|nm|μm|um|m)3' + _SE), '体积单位：如 cm3 → cm^3^（上标指数）'),
    # 幂：仅匹配带字面 ^ 的 10^n（作者手写乘方），裸数字不碰
    (re.compile(_SS + r'10\^(\d+)' + _SE), '幂次：10^6 → 10^6^（markdown 上标）'),
]

# 已带 markdown 上下标或 HTML sup/sub 的片段：剥离后再扫裸写，避免误报
_SUBSUP_WRAPPED_RE = re.compile(
    r'\^[^\^\s]+\^|~[^~\s]+~|<sup>.*?</sup>|<sub>.*?</sub>',
    re.IGNORECASE | re.DOTALL,
)

# F1：中文高置信错别字（保守只收确定性书写错字；的/地/得等主观字一律不收）。
_CHINESE_TYPOS = {
    '帐号': '账号', '登陆': '登录', '既使': '即使', '按装': '安装',
    '做为': '作为', '拌随': '伴随', '迫不急待': '迫不及待',
    '再接再励': '再接再厉', '一如继往': '一如既往', '言简意骇': '言简意赅',
    '甘败下风': '甘拜下风', '察颜观色': '察言观色', '脉博': '脉搏',
    '松驰': '松弛', '竭泽而鱼': '竭泽而渔', '金榜提名': '金榜题名',
    '美仑美奂': '美轮美奂', '病入膏盲': '病入膏肓', '穿流不息': '川流不息',
}

# F2：英文高置信拼写错误（固定错拼表，键=错拼，值=正确）。
# 只收"永不是合法英文词、也不像基因符号/缩写/化学式"的铁错拼，误报率≈0，故 ERROR 硬拦。
# 与 F1 中文错别字的区别：中文"登陆/做为"个别语境合法（登陆作战），F1 维持 WARN；
# occured/recieve 这类在英文里从不成立，可硬拦。（与 nsfc humanizer_zh 同表，保持一致）
_ENGLISH_MISSPELLINGS = {
    "occured": "occurred", "occuring": "occurring", "occurance": "occurrence",
    "occurence": "occurrence", "recieve": "receive", "recieved": "received",
    "seperate": "separate", "seperately": "separately", "definately": "definitely",
    "neccessary": "necessary", "accomodate": "accommodate", "enviroment": "environment",
    "existance": "existence", "acheive": "achieve", "acheived": "achieved",
    "arguement": "argument", "begining": "beginning", "beleive": "believe",
    "comparision": "comparison", "consistant": "consistent", "goverment": "government",
    "paralell": "parallel", "persistant": "persistent", "prefered": "preferred",
    "refered": "referred", "relevent": "relevant", "succesful": "successful",
    "successfull": "successful", "untill": "until", "writen": "written",
}
# ASCII 边界：汉字紧邻也抓（"细胞occured了"），不匹配更长英文词内部（reoccured 不误报），大小写不敏感。
_EN_MISSPELL_RE = re.compile(
    _SS + r'(' + '|'.join(re.escape(w) for w in _ENGLISH_MISSPELLINGS) + r')' + _SE,
    re.IGNORECASE,
)


def check_char_level(doc):
    """字符级检查。severity 级别：D1 半角标点 / D2 上下标裸写 / F1 中文错别字 = WARN；F2 英文拼写 = ERROR。

    注意：severity 级别 ≠ 是否硬阻断。D1(halfwidth_punct_in_cn) / D2(subsup_bare) / F2(english_misspelling)
    三项都是退出码硬门——主流程按 issue_summary 计数判定，任一 >0 即 check_quality.py 非零退出，
    与其 warning/error 标签无关（详见 __main__ 退出码逻辑，与 dod_checklist.json S9 一致）。
    唯 F1 中文错别字仅 WARN、不进硬门（含登陆等同形异义，硬拦会误伤）。

    跳过参考文献与摘要区（与 check_writing_style 同口径），避免英文/DOI 误报。
    例外：F2 英文拼写在摘要区也扫——英文主要出现在英文摘要，固定错拼表误报率≈0。
    """
    issues = []
    in_references = False
    in_abstract = False

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        lvl = heading_level(getattr(para.style, "name", ""))
        if lvl is not None:
            cls = classify_heading(text)
            if cls == "references":
                in_references = True
            elif in_references:
                in_references = False
            in_abstract = "摘要" in text or "abstract" in text.lower()
            continue

        if in_references:
            continue

        # F2：英文高置信拼写错误（ERROR 硬拦）。摘要区也扫——英文主要出现在英文摘要。
        for m in _EN_MISSPELL_RE.finditer(text):
            wrong = m.group(0)
            right = _ENGLISH_MISSPELLINGS[wrong.lower()]
            issues.append({
                'level': 'error',
                'category': '字符级-英文拼写',
                'message': f'英文拼写错误：{wrong} → {right}',
                'suggestion': f'改为"{right}"',
                'code': 'english_misspelling',
            })

        if in_abstract:
            continue

        # D1：中文句内夹半角标点
        for pat, half, full, desc in _HALFWIDTH_IN_CN:
            for m in pat.finditer(text):
                issues.append({
                    'level': 'warning',
                    'category': '字符级-半角标点',
                    'message': f'{desc}：…{m.group(0)}…（"{half}" → "{full}"）',
                    'suggestion': '中文句内标点应使用全角',
                    'code': 'halfwidth_punct_in_cn',
                })

        # D2：应上下标却裸写。检测输入用"非上下标 run 的拼接"——已被 markdown_to_docx
        # 渲染成 subscript/superscript run 的字符（如 CO₂ 的 "2"）不进检测文本，避免与
        # 导出器自相矛盾误报；真裸写（全 plain run）照进照报。仍剥离残留的 md/HTML 文本标记。
        subsup_source = "".join(
            r.text for r in para.runs
            if not r.font.superscript and not r.font.subscript
        )
        stripped = _SUBSUP_WRAPPED_RE.sub(' ', subsup_source)
        seen = set()
        for pat, desc in _SUBSUP_PATTERNS:
            for m in pat.finditer(stripped):
                key = (m.group(0), m.start())
                if key in seen:
                    continue
                seen.add(key)
                issues.append({
                    'level': 'warning',
                    'category': '字符级-上下标',
                    'message': f'疑似应上下标却裸写：{m.group(0)}（{desc}）',
                    'suggestion': '用 markdown ~下标~ / ^上标^ 标注',
                    'code': 'subsup_bare',
                })

        # F1：中文高置信错别字
        for wrong, right in _CHINESE_TYPOS.items():
            for m in re.finditer(re.escape(wrong), text):
                issues.append({
                    'level': 'warning',
                    'category': '字符级-错别字',
                    'message': f'疑似错别字：{wrong} → {right}',
                    'suggestion': f'确认是否应为"{right}"',
                    'code': 'chinese_typo',
                })

    return issues


def check_writing_style(doc):
    """
    检测写作风格违规：
    - 破折号（——）
    - 问句（正文应全部陈述）
    - 比喻修辞（如同、犹如、像...一样、...的桥梁等）
    - 主观夸大表述
    - 过度书面化/生僻词
    - 排比句式
    - scare quotes（引号包裹普通中文短语）
    - 解释性冒号（"概念：解释"装饰句式）
    """
    issues = []
    in_references = False
    in_abstract = False

    body_sentences = []  # 收集正文句子用于排比检测

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        lvl = heading_level(getattr(para.style, "name", ""))
        if lvl is not None:
            cls = classify_heading(text)
            if cls == "references":
                in_references = True
            elif in_references:
                in_references = False
            in_abstract = "摘要" in text or "abstract" in text.lower()
            continue

        # 跳过参考文献和摘要
        if in_references or in_abstract:
            continue

        # 逐条规则检测
        for pattern, level, category, message, suggestion in _STYLE_CHECKS:
            for m in pattern.finditer(text):
                start = max(0, m.start() - 8)
                end = min(len(text), m.end() + 8)
                snippet = text[start:end]
                issues.append({
                    'level': level,
                    'category': category,
                    'message': f'{message}：...{snippet}...',
                    'suggestion': suggestion,
                })

        # ── 中文 AI 禁词检测 ──────────────────────────────────────────────
        for pattern, category, message, suggestion in _CN_AI_BANNED:
            for m in pattern.finditer(text):
                start = max(0, m.start() - 8)
                end = min(len(text), m.end() + 8)
                snippet = text[start:end]
                issues.append({
                    'level': 'warning',
                    'category': category,
                    'message': f'{message}：...{snippet}...',
                    'suggestion': suggestion,
                })

        # ── 翻译腔软检测（全 info，不阻断）────────────────────────────────
        # 1) 名词化/虚化动词：进行/加以/予以/使得/得以
        for m in _TE_NOMINALIZE.finditer(text):
            start = max(0, m.start() - 8)
            end = min(len(text), m.end() + 8)
            issues.append({
                'level': 'info',
                'category': '翻译腔',
                'message': f'名词化/虚化动词（疑似翻译腔）：...{text[start:end]}...',
                'suggestion': '改为直接的动作动词，如"进行分析"→"分析"、"使得X提高"→"提高了X"',
                'code': 'translationese_nominalize',
            })
        # 2) 被字句密度：单段"被"过多
        bei_count = text.count('被')
        if bei_count >= _TE_BEI_PER_PARA:
            issues.append({
                'level': 'info',
                'category': '翻译腔',
                'message': f'被动句偏多（本段 {bei_count} 处"被"，疑似英式被动直译）',
                'suggestion': '中文优先主动句，仅在施事不明或强调受事时保留被动',
                'code': 'translationese_passive',
            })
        # 3) 长定语链：…的…的…的…的
        for m in _TE_LONG_ATTR_CHAIN.finditer(text):
            snippet = m.group(0)[:40] + ('…' if len(m.group(0)) > 40 else '')
            issues.append({
                'level': 'info',
                'category': '翻译腔',
                'message': f'超长前置定语链（疑似翻译腔）：{snippet}',
                'suggestion': '拆分多层"的"字定语，改用短句或后置从句',
                'code': 'translationese_long_attr',
            })

        # ── 中文句长检测 ──────────────────────────────────────────────────
        # 按中文句末标点（。！？）切句；空句跳过；只数中文字符
        cn_sentences_in_para = re.split(r'[。！？]', text)
        cn_lens = []
        for s in cn_sentences_in_para:
            s = s.strip()
            if not s:
                continue
            cn_char_count = len(re.findall(r'[一-鿿]', s))
            if cn_char_count < 4:
                continue  # 忽略过短残句（如引用编号后的残余）
            cn_lens.append(cn_char_count)
            if cn_char_count > _CN_SENT_MAX_CHARS:
                # 截取snippet（最多显示40字）
                snippet = s[:40] + ('…' if len(s) > 40 else '')
                issues.append({
                    'level': 'warning',
                    'category': '句长规范',
                    'message': f'中文单句过长（{cn_char_count} 字 > {_CN_SENT_MAX_CHARS} 字上限）：{snippet}',
                    'suggestion': '拆分长句：每句 ≤50 中文字符，从句嵌套 ≤2 层',
                    'code': 'cn_sentence_too_long',
                })
        # 连续3句长度变化 <5 字（单调句式告警）
        if len(cn_lens) >= 3:
            for k in range(len(cn_lens) - 2):
                trio = cn_lens[k:k+3]
                if max(trio) - min(trio) < _CN_SENT_MONOTONE_DIFF:
                    issues.append({
                        'level': 'info',
                        'category': '句长规范',
                        'message': (
                            f'连续 3 句字数差异 <{_CN_SENT_MONOTONE_DIFF} 字（{trio}），'
                            '句式单调，建议短句（≤15字）与长句（30-50字）交替'
                        ),
                        'suggestion': '调整句子长度节奏：短句(<15字)与长句(30-50字)交替出现',
                        'code': 'cn_sentence_monotone',
                    })
                    break  # 每段只报一次

        # 收集句子用于排比检测（按句号/分号切分）
        sentences = re.split(r'[。；;]', text)
        for s in sentences:
            s = s.strip()
            if len(s) >= 6:
                body_sentences.append(s)

    # 排比检测：连续句子以相同前缀开头
    if len(body_sentences) >= _PARALLELISM_MIN_COUNT:
        i = 0
        while i < len(body_sentences):
            # 取前4个字作为模式
            prefix = body_sentences[i][:4]
            if not prefix:
                i += 1
                continue
            run = 1
            j = i + 1
            while j < len(body_sentences) and body_sentences[j][:4] == prefix:
                run += 1
                j += 1
            if run >= _PARALLELISM_MIN_COUNT:
                examples = '；'.join(body_sentences[i:i+3])
                issues.append({
                    'level': 'warning',
                    'category': '修辞规范',
                    'message': f'疑似排比句式（连续{run}句以"{prefix}"开头）：{examples}...',
                    'suggestion': '避免重复句式结构，改用多样化的表达方式',
                })
                i = j
            else:
                i += 1

    return issues


def check_full_thesis_structure(doc, min_chapters=5):
    """
    全文结构门禁：
    1) 章节数不少于 min_chapters
    2) 第一章应为绪论/引言
    3) 最后一章应为总结/结论/展望类独立章节
    """
    issues = []
    chapter_titles = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        lvl = heading_level(getattr(para.style, "name", ""))
        if lvl is None:
            continue
        if is_chapter_heading_text(text):
            chapter_titles.append(text)

    min_chapters = max(1, int(min_chapters or 5))
    if len(chapter_titles) < min_chapters:
        issues.append({
            'level': 'error',
            'category': '结构',
            'message': f'章节数不足：{len(chapter_titles)} / {min_chapters}',
            'suggestion': f'需满足“独立绪论 + 多个研究章 + 独立总结章”，总章节数不少于 {min_chapters}'
        })
        return issues

    first_title = chapter_titles[0]
    first_norm = normalize_text(first_title)
    if ("绪论" not in first_norm) and ("引言" not in first_norm):
        issues.append({
            'level': 'error',
            'category': '结构',
            'message': f'第一章标题不符合绪论要求：{first_title}',
            'suggestion': '第一章应设置为“绪论”或“引言”性质章节'
        })

    last_title = chapter_titles[-1]
    last_norm = normalize_text(last_title)
    if not any(k in last_norm for k in ("结论", "总结", "小结", "展望")):
        issues.append({
            'level': 'error',
            'category': '结构',
            'message': f'最后一章标题不符合独立总结章要求：{last_title}',
            'suggestion': '最后一章应为独立总结/结论/展望章节'
        })

    if len(chapter_titles) >= 2:
        research_chapter_count = len(chapter_titles) - 2
        if research_chapter_count <= 0:
            issues.append({
                'level': 'error',
                'category': '结构',
                'message': '研究章节数量不足',
                'suggestion': '应保留“独立绪论章 + 多个研究章 + 独立总结章”组织形式'
            })

    return issues


def check_cross_chapter_coherence(doc, project_root=None):
    """章间衔接检查（软门禁，全 warning/info，不阻断）——防"订书机论文"。

    以 project_state.json 的 outline（研究主线表）为基准，核对：
      1) 绪论是否预告了每个研究章的贡献（研究章标题应在绪论正文出现）；
      2) 结论是否逐章综合（研究章标题应在结论正文出现）；
      3) 每个研究章开头是否回指主线（首段含承接/主线线索词）。
    outline 缺失或为空则跳过（空 outline 由 Step 0.5 门禁另行拦截）。
    """
    issues = []
    if not project_root:
        return issues
    ps_path = os.path.join(project_root, "project_state.json")
    if not os.path.isfile(ps_path):
        return issues
    try:
        with open(ps_path, "r", encoding="utf-8") as f:
            ps = json.load(f)
    except Exception:
        return issues
    outline = ps.get("outline")
    if not isinstance(outline, list) or not outline:
        return issues

    # 研究章：outline 中带 core_argument 的条目（绪论/结论一般不写 core_argument）
    research_entries = [
        e for e in outline
        if isinstance(e, dict) and str(e.get("core_argument", "")).strip()
        and str(e.get("title", "")).strip()
    ]
    if not research_entries:
        return issues

    # 将 docx 切分为章块：{chapter_num: {"title":..., "body":[首段..], "text": 全文}}
    chapters = {}
    cur = None
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        lvl = heading_level(getattr(para.style, "name", ""))
        if lvl == 1 and is_chapter_heading_text(text):
            m = re.search(r'第\s*(\d+)\s*章', text)
            cur = int(m.group(1)) if m else None
            if cur is not None:
                chapters[cur] = {"title": text, "body_paras": [], "text": ""}
            continue
        if cur is not None and cur in chapters:
            if lvl is None:  # 正文段
                chapters[cur]["body_paras"].append(text)
            chapters[cur]["text"] += text

    # 绪论/结论章文本
    intro_text = ""
    concl_text = ""
    for num, blk in chapters.items():
        norm = normalize_text(blk["title"])
        if ("绪论" in norm or "引言" in norm) and not intro_text:
            intro_text = blk["text"]
        if any(k in norm for k in ("结论", "总结", "小结", "展望")):
            concl_text = blk["text"]  # 取最后一个匹配（结论通常靠后）

    for e in research_entries:
        anchor = str(e.get("title", "")).strip()
        ch = e.get("chapter")
        label = f"研究章「{anchor}」" + (f"（第{ch}章）" if ch is not None else "")
        if intro_text and anchor not in intro_text:
            issues.append({
                'level': 'warning',
                'category': '章间衔接',
                'message': f'绪论未预告{label}的贡献点',
                'suggestion': '在绪论"研究内容/技术路线"处逐一预告各研究章要解决的问题与贡献',
            })
        if concl_text and anchor not in concl_text:
            issues.append({
                'level': 'warning',
                'category': '章间衔接',
                'message': f'结论未跨章综合{label}',
                'suggestion': '在结论章逐章回收各研究章的核心发现并综合成整体结论',
            })
        # 研究章开头回指主线
        if ch is not None and ch in chapters:
            head = "".join(chapters[ch]["body_paras"][:2])
            if head and not re.search(r'上一章|前一章|前文|前述|在.{0,12}基础上|本章|围绕|针对|承接', head):
                issues.append({
                    'level': 'info',
                    'category': '章间衔接',
                    'message': f'{label}开头缺少与研究主线的承接语',
                    'suggestion': '研究章首段先回指上一章结论或全文科学问题，再引出本章目标',
                })

    return issues


def check_section_order(doc):
    """
    全文结构顺序校验：验证各部分出现顺序是否符合规范。

    期望顺序（允许缺省，但不允许乱序）：
    封面 → 独创性声明 → 摘要 → 英文摘要 → 目录 → 缩略语表 →
    正文(第1章..第N章) → 参考文献 → 致谢 → 攻读期间成果 → 附录
    """
    issues = []

    # 定义有序标签
    _ORDERED_LABELS = [
        ('封面', ['封面', '题名']),
        ('独创性声明', ['独创性', '授权']),
        ('中文摘要', ['摘要']),
        ('英文摘要', ['abstract']),
        ('目录', ['目录']),
        ('缩略语表', ['缩略', '符号']),
        ('正文', []),          # 特殊：匹配 "第X章"
        ('参考文献', ['参考文献']),
        ('致谢', ['致谢']),
        ('攻读期间成果', ['攻读', '成果']),
        ('附录', ['附录']),
    ]

    def _match_label(text):
        t = text.strip().lower()
        if re.search(r'第\d+章', t):
            return '正文'
        for label, keywords in _ORDERED_LABELS:
            for kw in keywords:
                if kw in t:
                    return label
        return None

    # 收集 H1 标题的标签序列（去重连续相同标签）
    seen_labels = []
    for para in doc.paragraphs:
        lvl = heading_level(getattr(para.style, 'name', ''))
        if lvl != 1:
            continue
        label = _match_label(para.text)
        if label is None:
            continue
        if not seen_labels or seen_labels[-1] != label:
            seen_labels.append(label)

    # 检查顺序：seen_labels 应该是 _ORDERED_LABELS 标签的子序列
    label_order = [lbl for lbl, _ in _ORDERED_LABELS]
    expected_idx = 0
    for actual_label in seen_labels:
        # 在 label_order 中向前查找
        found = False
        for i in range(expected_idx, len(label_order)):
            if label_order[i] == actual_label:
                expected_idx = i + 1
                found = True
                break
        if not found:
            issues.append({
                'level': 'error',
                'category': '结构顺序',
                'message': f'"{actual_label}" 出现位置不符合规范顺序',
                'suggestion': f'期望顺序：封面→独创性声明→摘要→英文摘要→目录→缩略语表→正文→参考文献→致谢→成果→附录'
            })

    return issues


def check_page_breaks_between_chapters(doc):
    """
    检查章节之间是否存在分页符或分节符。

    扫描所有 Heading 1 段落，检查其前一个段落是否包含分页符
    （w:br type="page"）或分节符（w:sectPr type="nextPage"）。
    跳过文档第一个 H1（无需前置分页）。
    """
    issues = []

    h1_paras = []
    for para in doc.paragraphs:
        lvl = heading_level(getattr(para.style, 'name', ''))
        if lvl == 1:
            h1_paras.append(para)

    for idx, para in enumerate(h1_paras):
        if idx == 0:
            continue  # 第一个 H1 不需要前置分页

        p_elem = para._element
        prev = p_elem.getprevious()

        has_break = False

        if prev is not None:
            # 检查前一段落是否有分页符 <w:br w:type="page"/>
            for br in prev.iter(qn('w:br')):
                if br.get(qn('w:type')) == 'page':
                    has_break = True
                    break

            # 检查前一段落是否有分节符 <w:sectPr>
            if not has_break:
                pPr = prev.find(qn('w:pPr'))
                if pPr is not None and pPr.find(qn('w:sectPr')) is not None:
                    has_break = True

        # 也检查 H1 段落自身是否设置了 pageBreakBefore
        if not has_break:
            pPr = p_elem.find(qn('w:pPr'))
            if pPr is not None:
                pgBB = pPr.find(qn('w:pageBreakBefore'))
                if pgBB is not None and pgBB.get(qn('w:val'), 'true') != 'false':
                    has_break = True

        if not has_break:
            issues.append({
                'level': 'warning',
                'category': '分页',
                'message': f'章节标题 "{para.text.strip()}" 前缺少分页符或分节符',
                'suggestion': '每章应从新页开始，请在章标题前插入分页符'
            })

    return issues


# ---------------------------------------------------------------------------
# 三线表格式检测
# ---------------------------------------------------------------------------

def _get_cell_border(cell, edge):
    """
    获取单元格某条边框的属性。
    返回 dict: {"val": str, "sz": int} 或 None（无边框定义）。
    """
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        return None
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        return None
    el = tcBorders.find(qn(f'w:{edge}'))
    if el is None:
        return None
    val = el.get(qn('w:val'))
    sz_str = el.get(qn('w:sz'))
    sz = int(sz_str) if sz_str else 0
    return {"val": val, "sz": sz}


def check_table_format(doc):
    """
    检查 Word 文档中的表格是否符合三线表规范：
    - 顶部边框 1.5pt (sz=12)
    - 底部边框 1.5pt (sz=12)
    - 表头分隔线 0.5pt (sz=4)
    - 无竖线
    """
    issues = []

    for t_idx, table in enumerate(doc.tables):
        num_rows = len(table.rows)
        if num_rows < 2:
            continue  # 单行表格跳过

        table_label = f'表格 {t_idx + 1}'

        # 检查第一行顶部边框
        for cell in table.rows[0].cells:
            border = _get_cell_border(cell, 'top')
            if border is None or border['val'] == 'none':
                issues.append({
                    'level': 'error',
                    'category': '三线表',
                    'message': f'{table_label}：首行缺少顶部边框',
                    'suggestion': '三线表首行顶部应有 1.5pt 实线边框'
                })
                break
            elif border['sz'] < 10:  # 允许 10-14 范围（约 1.25-1.75pt）
                issues.append({
                    'level': 'warning',
                    'category': '三线表',
                    'message': f'{table_label}：首行顶部边框过细（{border["sz"]/8:.1f}pt）',
                    'suggestion': '三线表顶部边框应为 1.5pt'
                })
                break

        # 检查最后一行底部边框
        for cell in table.rows[-1].cells:
            border = _get_cell_border(cell, 'bottom')
            if border is None or border['val'] == 'none':
                issues.append({
                    'level': 'error',
                    'category': '三线表',
                    'message': f'{table_label}：末行缺少底部边框',
                    'suggestion': '三线表末行底部应有 1.5pt 实线边框'
                })
                break
            elif border['sz'] < 10:
                issues.append({
                    'level': 'warning',
                    'category': '三线表',
                    'message': f'{table_label}：末行底部边框过细（{border["sz"]/8:.1f}pt）',
                    'suggestion': '三线表底部边框应为 1.5pt'
                })
                break

        # 检查表头分隔线（第一行底部，应为 0.75pt / sz=6）
        for cell in table.rows[0].cells:
            border = _get_cell_border(cell, 'bottom')
            if border is None or border['val'] == 'none':
                issues.append({
                    'level': 'error',
                    'category': '三线表',
                    'message': f'{table_label}：缺少表头分隔线',
                    'suggestion': '三线表表头与表体之间应有 0.75pt 分隔线'
                })
                break
            elif border['sz'] > 8:  # 允许 sz 2-8（0.25-1.0pt），超过则过粗
                issues.append({
                    'level': 'warning',
                    'category': '三线表',
                    'message': f'{table_label}：表头分隔线过粗（{border["sz"]/8:.2f}pt，应为 0.75pt/sz=6）',
                    'suggestion': '三线表表头分隔线应为 0.75pt（sz=6）'
                })
                break

        # 检查竖线（不应存在）
        has_vertical = False
        for row in table.rows:
            for cell in row.cells:
                for edge in ('left', 'right'):
                    border = _get_cell_border(cell, edge)
                    if border and border['val'] not in ('none', 'nil', None) and border['sz'] > 0:
                        has_vertical = True
                        break
                if has_vertical:
                    break
            if has_vertical:
                break
        if has_vertical:
            issues.append({
                'level': 'error',
                'category': '三线表',
                'message': f'{table_label}：存在竖线',
                'suggestion': '三线表不应有竖线，请移除所有垂直边框'
            })

    return issues


def check_abbreviation_consistency(doc, project_root=None):
    """
    检查缩略语一致性：
    1. 同一缩略语是否在多个章节重复展开
    2. 正文中出现的缩略语是否都在注册表中
    """
    issues = []

    if load_registry is None or extract_abbreviations is None:
        return issues

    # 收集全文文本，按章节分组
    chapter_texts = {}
    current_chapter = "0"
    current_text_lines = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        lvl = heading_level(getattr(para.style, "name", ""))
        if lvl is not None and is_chapter_heading_text(text):
            if current_text_lines:
                chapter_texts[current_chapter] = "\n".join(current_text_lines)
            m = re.search(r'[一二三四五六七八九十百千万零〇0-9]+', text)
            current_chapter = m.group(0) if m else current_chapter
            current_text_lines = []
            continue
        current_text_lines.append(text)

    if current_text_lines:
        chapter_texts[current_chapter] = "\n".join(current_text_lines)

    # 提取每章中出现的缩略语展开
    abbr_first_seen = {}  # abbr -> first chapter
    for ch_key in sorted(chapter_texts.keys(), key=lambda x: str(x)):
        ch_text = chapter_texts[ch_key]
        extracted = extract_abbreviations(ch_text)
        for item in extracted:
            abbr = item["abbr"]
            if abbr not in abbr_first_seen:
                abbr_first_seen[abbr] = ch_key
            else:
                if abbr_first_seen[abbr] != ch_key:
                    issues.append({
                        'level': 'warning',
                        'category': '缩略语',
                        'message': f'缩略语 {abbr} 在第{ch_key}章重复展开（首次出现于第{abbr_first_seen[abbr]}章）',
                        'suggestion': f'第{ch_key}章中应直接使用 {abbr}，无需再次展开全称',
                    })

    # 如果有注册表，检查正文中的大写缩写是否都已注册
    if project_root:
        try:
            registry = load_registry(project_root)
        except Exception:
            registry = {}
        if registry:
            full_text = "\n".join(chapter_texts.values())
            # 匹配独立的大写缩写词（2-15字符）
            standalone_abbrs = set(re.findall(r'\b([A-Z][A-Za-z0-9\-]{1,14})\b', full_text))
            # 过滤常见非缩略语词
            common_words = {
                "The", "In", "On", "At", "For", "And", "But", "Or", "Not", "No",
                "Yes", "OK", "DNA", "RNA", "pH", "UV", "IR", "NMR", "MS", "GC",
                "HPLC", "SEM", "TEM", "XRD", "XPS", "BET", "DLS", "TGA", "DSC",
                "IC50", "EC50", "LD50", "ED50",
            }
            # 英文标题词白名单（首字母大写的普通英语词，非学术缩略语）
            _TITLE_WORD_RE = re.compile(
                r'^(?:A|An|The|Is|Are|Was|Were|Be|Has|Have|Had|'
                r'This|That|These|Those|'
                r'Of|In|On|At|For|And|But|Or|Not|No|Nor|'
                r'Its|Their|Our|His|Her|'
                r'By|To|With|From|Into|Than|After|Before|Between|'
                r'Cancer|Tumor|Tumour|Cell|Cells|Gene|Genes|Protein|Proteins|'
                r'Study|Role|Effect|Effects|Impact|Analysis|Review|'
                r'Carcinoma|Sarcoma|Lymphoma|Leukemia|'
                r'Expression|Regulation|Activation|Inhibition|'
                r'Treatment|Therapy|Mechanism|Pathway|Signaling|'
                r'Human|Mouse|Rat|'
                r'Figure|Table|Fig|[A-Z])$'
            )
            for abbr in standalone_abbrs:
                if len(abbr) < 2:
                    continue
                if abbr in common_words:
                    continue
                # 跳过英文标题词
                if _TITLE_WORD_RE.match(abbr):
                    continue
                # 跳过含连字符的化合物名（如 YKL-5-、SB-4-等药物/化合物编号）
                if '-' in abbr:
                    continue
                # 跳过氨基酸位点标记（如 Ser5 Thr7）：字母后跟数字
                if re.match(r'^[A-Z][a-z]{0,3}\d+$', abbr):
                    continue
                if abbr not in registry and abbr.upper() not in {k.upper() for k in registry}:
                    # 只报告出现3次以上的未注册缩写
                    count = len(re.findall(r'\b' + re.escape(abbr) + r'\b', full_text))
                    if count >= 3:
                        issues.append({
                            'level': 'info',
                            'category': '缩略语',
                            'message': f'缩写 {abbr} 出现 {count} 次但未在缩略语注册表中',
                            'suggestion': f'如为专业缩略语，建议注册到 abbreviation_registry.json',
                        })

    # 交叉引用验证：注册表中的 first_chapter/first_section 是否与 markdown 文件一致
    if project_root and validate_cross_references is not None:
        try:
            xref_result = validate_cross_references(project_root)
            for detail in xref_result.get("details", []):
                if detail.get("status") == "invalid":
                    issues.append({
                        'level': 'warning',
                        'category': '缩略语交叉引用',
                        'message': f'缩略语 {detail["abbr"]} 交叉引用失败：{detail.get("reason", "unknown")}',
                        'suggestion': '请检查注册表中 first_chapter/first_section 是否正确，或对应 markdown 文件是否包含该缩略语展开',
                    })
        except Exception:
            pass

    return issues


def check_paragraph_formatting(doc):
    """检查段落格式是否规范"""
    issues = []
    
    for i, para in enumerate(doc.paragraphs):
        if not para.text.strip():
            continue
        
        # 跳过标题
        if heading_level(getattr(para.style, "name", "")) is not None:
            continue
        
        # 检查首行缩进（正文应有首行缩进）
        if para.style.name == 'Normal':
            if not para.paragraph_format.first_line_indent:
                issues.append({
                    'level': 'info',
                    'category': '格式',
                    'location': f'第 {i+1} 段',
                    'message': '正文段落缺少首行缩进',
                    'suggestion': '正文段落应设置首行缩进 2 字符'
                })
        
        # 检查行距（应为 20 磅）
        spacing_pt = line_spacing_pt(para)
        if spacing_pt is not None:
            if abs(spacing_pt - 20.0) > 0.5:
                issues.append({
                    'level': 'info',
                    'category': '格式',
                    'location': f'第 {i+1} 段',
                    'message': f'行距不符合要求：{spacing_pt:.2f} 磅',
                    'suggestion': '正文行距应设置为固定值 20 磅'
                })
    
    return issues


def _get_east_asian_font(run):
    """从 run 的 XML 中提取东亚字体名称。"""
    rPr = run._element.find(qn('w:rPr'))
    if rPr is None:
        return None
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        return None
    return rFonts.get(qn('w:eastAsia'))


def _align_value(alignment_name):
    mapping = {
        "left": 0,
        "center": 1,
        "right": 2,
        "justify": 3,
    }
    return mapping.get(str(alignment_name or "").strip().lower())


def _style_spec_to_quality_spec(style_spec, label, match):
    spec = style_spec if isinstance(style_spec, dict) else {}
    first_indent_cm = spec.get("first_line_indent_cm")
    indent_tol = 30000
    quality_spec = {
        "match": match,
        "label": label,
        "font_size_pt": spec.get("font_size_pt"),
        "bold": spec.get("bold"),
        "alignment": _align_value(spec.get("alignment")),
        "line_spacing_pt": spec.get("line_spacing_pt") if spec.get("line_spacing_rule") == "exact" else None,
        "line_spacing_rule": spec.get("line_spacing_rule"),
        "space_before_pt": spec.get("space_before_pt"),
        "space_after_pt": spec.get("space_after_pt"),
        "font_latin": spec.get("font_latin"),
        "font_east_asia": spec.get("font_east_asia"),
    }
    if first_indent_cm is not None:
        emu = int(first_indent_cm * 360000)
        quality_spec["first_indent_emu_min"] = emu
        quality_spec["first_indent_emu_max"] = emu
        quality_spec["indent_tol"] = indent_tol
    return quality_spec


def check_word_format_compliance(doc, format_context=None):
    """
    全面检查 Word 文档格式是否符合内置默认模板规范。
    检查页面布局、字体、字号、行距、缩进、段前段后间距等。
    """
    issues = []
    EMU_PER_CM = 360000

    # ── 1. 页面布局检查 ──
    format_context = format_context or build_format_render_context({})
    page_margins = format_context.get("page_margins_cm", {})
    expected_page = {
        'width_cm': 21.0,
        'height_cm': 29.7,
        'top_cm': page_margins.get("top", 2.54),
        'bottom_cm': page_margins.get("bottom", 2.54),
        'left_cm': page_margins.get("left", 3.17),
        'right_cm': page_margins.get("right", 3.17),
    }
    page_tol = 0.2
    margin_tol = 0.1

    for idx, section in enumerate(doc.sections):
        sec_label = f'第 {idx + 1} 节'

        if hasattr(section, 'page_width') and section.page_width is not None:
            w_cm = section.page_width / EMU_PER_CM
            if abs(w_cm - expected_page['width_cm']) > page_tol:
                issues.append({
                    'level': 'error', 'category': '页面布局',
                    'message': f'{sec_label}纸张宽度不符：{w_cm:.2f}cm（应为 21.0cm）',
                    'suggestion': '请将纸张大小设置为 A4（21.0×29.7cm）'
                })
                break

        if hasattr(section, 'page_height') and section.page_height is not None:
            h_cm = section.page_height / EMU_PER_CM
            if abs(h_cm - expected_page['height_cm']) > page_tol:
                issues.append({
                    'level': 'error', 'category': '页面布局',
                    'message': f'{sec_label}纸张高度不符：{h_cm:.2f}cm（应为 29.7cm）',
                    'suggestion': '请将纸张大小设置为 A4（21.0×29.7cm）'
                })
                break

        margin_checks = [
            ('top_margin', 'top_cm', '上边距'),
            ('bottom_margin', 'bottom_cm', '下边距'),
            ('left_margin', 'left_cm', '左边距'),
            ('right_margin', 'right_cm', '右边距'),
        ]
        for attr, key, label in margin_checks:
            if hasattr(section, attr) and getattr(section, attr) is not None:
                val_cm = getattr(section, attr) / EMU_PER_CM
                if abs(val_cm - expected_page[key]) > margin_tol:
                    issues.append({
                        'level': 'error', 'category': '页面布局',
                        'message': f'{sec_label}{label}不符：{val_cm:.2f}cm（应为 {expected_page[key]}cm）',
                        'suggestion': f'请将{label}设置为 {expected_page[key]}cm'
                    })
        # 只报告第一个违规节
        if any(i['category'] == '页面布局' for i in issues):
            break

    # ── 2. 样式规范定义 ──
    style_profile = format_context.get("style_profile", {}) if isinstance(format_context, dict) else {}
    style_specs = {
        'heading1': _style_spec_to_quality_spec(
            style_profile.get("heading1", {}),
            '一级标题',
            lambda s: s and ('Heading 1' in s or '标题 1' in s),
        ),
        'heading2': _style_spec_to_quality_spec(
            style_profile.get("heading2", {}),
            '二级标题',
            lambda s: s and ('Heading 2' in s or '标题 2' in s),
        ),
        'heading3': _style_spec_to_quality_spec(
            style_profile.get("heading3", {}),
            '三级标题',
            lambda s: s and ('Heading 3' in s or '标题 3' in s),
        ),
        'normal': _style_spec_to_quality_spec(
            style_profile.get("body", {}),
            '正文',
            lambda s: s and (s == 'Normal' or s == '正文'),
        ),
    }

    FONT_SIZE_TOL = 0.5
    SPACING_TOL = 1.0

    # 收集每种样式前 5 个段落
    style_samples = {k: [] for k in style_specs}
    for para in doc.paragraphs:
        style_name = para.style.name if para.style else None
        for key, spec in style_specs.items():
            if spec['match'](style_name) and len(style_samples[key]) < 5:
                style_samples[key].append(para)

    # 记录每个类别已报告的违规类型，避免重复
    reported = set()

    def _report(key, check_name, level, message, suggestion):
        tag = f'{key}:{check_name}'
        if tag not in reported:
            reported.add(tag)
            issues.append({
                'level': level, 'category': '格式规范',
                'message': message, 'suggestion': suggestion,
            })

    # ── 3. 逐样式检查 ──
    for key, spec in style_specs.items():
        label = spec['label']
        for para in style_samples[key]:
            # 字号检查
            if para.runs and spec['font_size_pt'] is not None:
                run = para.runs[0]
                if run.font.size is not None and hasattr(run.font.size, 'pt'):
                    actual = run.font.size.pt
                    if abs(actual - spec['font_size_pt']) > FONT_SIZE_TOL:
                        _report(key, 'font_size', 'error',
                                f'{label}字号不符：{actual}pt（应为 {spec["font_size_pt"]}pt）',
                                f'请将{label}字号设置为 {spec["font_size_pt"]}pt')

            # 加粗检查
            if para.runs and spec['bold'] is not None:
                run = para.runs[0]
                actual_bold = run.font.bold
                if actual_bold is not None and actual_bold != spec['bold']:
                    expected_str = '加粗' if spec['bold'] else '常规（非加粗）'
                    _report(key, 'bold', 'error',
                            f'{label}加粗属性不符（应为{expected_str}）',
                            f'请将{label}设置为{expected_str}')

            # 对齐方式检查
            if spec['alignment'] is not None:
                actual_align = para.paragraph_format.alignment
                if actual_align is not None:
                    align_val = actual_align if isinstance(actual_align, int) else actual_align.value if hasattr(actual_align, 'value') else None
                    if align_val is not None and align_val != spec['alignment']:
                        align_names = {0: '左对齐', 1: '居中', 2: '右对齐', 3: '两端对齐'}
                        expected_name = align_names.get(spec['alignment'], str(spec['alignment']))
                        _report(key, 'alignment', 'error',
                                f'{label}对齐方式不符（应为{expected_name}）',
                                f'请将{label}对齐方式设置为{expected_name}')

            # 行距检查
            if spec['line_spacing_pt'] is not None:
                actual_ls = line_spacing_pt(para)
                if actual_ls is not None and abs(actual_ls - spec['line_spacing_pt']) > SPACING_TOL:
                    _report(key, 'line_spacing', 'warning',
                            f'{label}行距不符：{actual_ls:.1f}pt（应为 {spec["line_spacing_pt"]}pt）',
                            f'请将{label}行距设置为固定值 {spec["line_spacing_pt"]}pt')

            # 段前间距检查
            if spec.get('space_before_pt') is not None:
                sb = para.paragraph_format.space_before
                if sb is not None and hasattr(sb, 'pt'):
                    actual_sb = sb.pt
                    if abs(actual_sb - spec['space_before_pt']) > SPACING_TOL:
                        _report(key, 'space_before', 'warning',
                                f'{label}段前间距不符：{actual_sb:.1f}pt（应为 {spec["space_before_pt"]}pt）',
                                f'请将{label}段前间距设置为 {spec["space_before_pt"]}pt')

            # 段后间距检查
            if spec.get('space_after_pt') is not None:
                sa = para.paragraph_format.space_after
                if sa is not None and hasattr(sa, 'pt'):
                    actual_sa = sa.pt
                    if abs(actual_sa - spec['space_after_pt']) > SPACING_TOL:
                        _report(key, 'space_after', 'warning',
                                f'{label}段后间距不符：{actual_sa:.1f}pt（应为 {spec["space_after_pt"]}pt）',
                                f'请将{label}段后间距设置为 {spec["space_after_pt"]}pt')

            # 首行缩进检查（仅正文）
            if spec.get('first_indent_emu_min') is not None:
                fi = para.paragraph_format.first_line_indent
                if fi is not None:
                    fi_val = int(fi)
                    emu_min = spec['first_indent_emu_min'] - spec.get('indent_tol', 0)
                    emu_max = spec['first_indent_emu_max'] + spec.get('indent_tol', 0)
                    if fi_val < emu_min or fi_val > emu_max:
                        fi_cm = fi_val / EMU_PER_CM
                        _report(key, 'first_indent', 'warning',
                                f'{label}首行缩进不符：{fi_cm:.2f}cm（应约为 0.74cm）',
                                '请将正文首行缩进设置为 2 字符（约 0.74cm）')

            # 字体名称检查（标题段落）
            if key.startswith('heading') and para.runs:
                run = para.runs[0]
                # 西文字体
                if run.font.name is not None and run.font.name != 'Times New Roman':
                    _report(key, 'font_latin', 'info',
                            f'{label}西文字体不符：{run.font.name}（应为 Times New Roman）',
                            f'请将{label}西文字体设置为 Times New Roman')
                # 东亚字体
                ea_font = _get_east_asian_font(run)
                if ea_font is not None:
                    expected_ea = spec.get('font_east_asia')
                    if expected_ea and ea_font != expected_ea:
                        _report(key, 'font_ea', 'info',
                                f'{label}中文字体不符：{ea_font}（应为 {expected_ea}）',
                                f'请将{label}中文字体设置为 {expected_ea}')

            # 正文字体名称检查
            if key == 'normal' and para.runs:
                run = para.runs[0]
                # 西文字体应为 Times New Roman
                expected_latin = spec.get('font_latin') or 'Times New Roman'
                if run.font.name is not None and run.font.name != expected_latin:
                    _report(key, 'font_latin', 'error',
                            f'{label}西文字体不符：{run.font.name}（应为 {expected_latin}）',
                            f'请将{label}西文字体设置为 {expected_latin}')
                # 中文字体
                ea_font = _get_east_asian_font(run)
                expected_ea = spec.get('font_east_asia') or 'SimSun'
                if ea_font is not None and ea_font != expected_ea:
                    _report(key, 'font_ea', 'error',
                            f'{label}中文字体不符：{ea_font}（应为 {expected_ea}）',
                            f'请将{label}中文字体设置为 {expected_ea}')

    return issues


# ---------------------------------------------------------------------------
# Markdown 文件质量检测
# ---------------------------------------------------------------------------

def check_markdown_quality(md_path, md_checks='all'):
    """
    对单个 Markdown 文件进行质量检测。
    返回 (issues_list, stats_dict)。

    md_checks='all'（默认，语义与历史一致）返回全部 13 类检查的 issue；
    md_checks='xref' 只返回 `交叉引用` 类 —— 全文总检要的是窄口：其余 12 类在生产
    从未跑过，`占位标记` 会逐条命中 sci2doc 自己强制的 [图]/[表]/[实验] 标记
    （30 图博论 −30 分），会把总分压穿 80 线、让退出码因与交叉引用无关的理由翻 1。
    stats 不受影响（统计照常全量）。
    """
    issues = []
    stats = {
        'chinese_chars': 0,
        'heading_count': 0,
        'citation_count': 0,
        'figure_count': 0,
        'table_count': 0,
    }

    if not os.path.isfile(md_path):
        return issues, stats

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    full_text = ''.join(lines)

    # ---- 1) 中文字数统计 ----
    stats['chinese_chars'] = len(re.findall(r'[\u4e00-\u9fff]', full_text))

    # ---- 辅助：代码块状态追踪 ----
    in_code_block = False
    heading_levels_seen = []  # (行号, 级别)

    fig_re = re.compile(r'图\s*(\d+)-(\d+)')
    tbl_re = re.compile(r'表\s*(\d+)-(\d+)')
    list_re = re.compile(r'^(\s*)([-*•]|\d+\.)\s+')

    fig_refs = []  # (chapter, seq)
    tbl_refs = []  # (chapter, seq)

    # ---- A5 交叉引用有效性：收集"目标存在集合" + "待核对引用" ----
    # 目标集合 = 全文实际出现过的编号（无论定义还是引用都计入，保守取并集）。
    fig_targets = set()     # "X-Y" 字符串
    tbl_targets = set()     # "X-Y"
    section_targets = set()  # 章节号字符串，如 "2" / "2.1" / "2.1.3"
    appendix_targets = set()  # 附录字母，如 "A"
    crossrefs = []  # (line_no, kind, target_str, raw_snippet)

    # 引导词表：中文博论高频五个（"参见/详见/另见"含"见"子串，长者在前保证 snippet 完整；
    # "参阅"不含"见"，是真正新增的引导词）
    _A5_LEAD = r'(?:参见|详见|另见|参阅|见)'
    _A5_FT_NUM = r'[0-9]+[-–][0-9]+'
    _A5_FT_ITEM = r'(?:图|表)\s*' + _A5_FT_NUM
    # 并列第二项及以后："见图2-1和图2-2" —— 必须整体纳入同一次匹配，
    # 否则减法构造的目标集合会把残留的"图2-2"当成定义（扩了等于没扩）。
    _A5_FT_TAIL = r'(?:\s*[和与、]\s*' + _A5_FT_ITEM + r')*'
    # 图/表引用统一扫描：引导词式 + "如…所示/所列/可见"式，两式都带并列尾巴。
    _A5_FT_REF_RE = re.compile(
        r'(?:' + _A5_LEAD + r'\s*' + _A5_FT_ITEM + _A5_FT_TAIL + r')'
        r'|(?:如\s*' + _A5_FT_ITEM + _A5_FT_TAIL + r'\s*(?:所示|所列|可见))'
    )
    # 从一段引用里逐项拆出 (图|表, 编号)，供并列项各自登记
    _A5_FT_ITEM_RE = re.compile(r'(图|表)\s*(' + _A5_FT_NUM + r')')
    _A5_SEC_REF_RE = re.compile(_A5_LEAD + r'\s*第?\s*([0-9]+(?:\.[0-9]+)*)\s*节')
    _A5_CHAP_REF_RE = re.compile(_A5_LEAD + r'\s*第\s*([0-9]+)\s*章')
    _A5_APP_REF_RE = re.compile(r'见?附录\s*([A-Z])')
    _A5_EN_SEC_RE = re.compile(r'Section\s+([0-9]+(?:\.[0-9]+)*)')
    _A5_EN_FIG_RE = re.compile(r'Figure\s+([0-9]+[-–][0-9]+)')
    _A5_EN_TBL_RE = re.compile(r'Table\s+([0-9]+[-–][0-9]+)')
    # 正向题注识别：行首题注（含 sci2doc 的 `[图]`/`[表]` 标记前缀）无条件算定义，
    # 不受任何引用正则剥离影响。缺了它，扩引用正则会把假阴翻成假阳。
    _A5_CAPTION_DEF_RE = re.compile(
        r'^\s*(?:\[图\]|\[表\])?\s*(图|表)\s*([0-9]+[-–][0-9]+)\s*[：:．.、]?'
    )
    # 章节标题号：行首 # 标题里的 "第N章" 或前导 "N.M" 数字串
    _A5_HEAD_CN_CHAP_RE = re.compile(r'^#{1,3}\s*第\s*([0-9]+)\s*章')
    # 放宽到 6 级、且编号后不强制空格：真稿常见 `#### 3.1.1.1 缓冲液`、`### 3.2.1材料来源`，
    # 旧式（^#{1,3} + 必须跟空格）登记不进目标集合 → 引用它们全被误报断链。
    # 本正则唯一去处是往 section_targets **加**元素，判定是集合差，故只会少报、不会多报；
    # 四级标题本身仍由下方"标题层级"error 单独拦下，信号不丢。
    # 副作用（可接受）：`## 2020年研究进展` 会把 "2020" 也登记成节号，仅当全文出现
    # "见2020节"/"见第2020章" 这类引用时才可能漏报断链——真稿中不存在这种编号。
    _A5_HEAD_NUM_RE = re.compile(r'^#{1,6}\s*([0-9]+(?:\.[0-9]+)*)')
    _A5_APP_HEAD_RE = re.compile(r'^#{1,3}\s*附录\s*([A-Z])')

    for line_no, raw_line in enumerate(lines, start=1):
        line = raw_line.rstrip('\n')

        # 代码块切换
        if line.lstrip().startswith('```'):
            in_code_block = not in_code_block
            continue

        if in_code_block:
            continue

        # ---- 2) 标题层级检查 ----
        heading_match = re.match(r'^(#{1,9})\s+', line)
        if heading_match:
            level = len(heading_match.group(1))
            stats['heading_count'] += 1
            heading_levels_seen.append((line_no, level))

            # A5：从标题登记章节号 / 附录号作为可引用目标
            m_chap = _A5_HEAD_CN_CHAP_RE.match(line)
            if m_chap:
                section_targets.add(m_chap.group(1))
            m_num = _A5_HEAD_NUM_RE.match(line)
            if m_num:
                num = m_num.group(1)
                section_targets.add(num)
                # "2.1.3" 也意味着第 2 章、2.1 节存在 → 登记所有前缀
                parts = num.split('.')
                for k in range(1, len(parts)):
                    section_targets.add('.'.join(parts[:k]))
            m_app = _A5_APP_HEAD_RE.match(line)
            if m_app:
                appendix_targets.add(m_app.group(1))

            if level > 3:
                issues.append({
                    'level': 'error',
                    'category': '标题层级',
                    'message': f'第 {line_no} 行：使用了 {level} 级标题（最多允许 3 级）',
                    'suggestion': '论文 Markdown 仅使用 #、##、### 三级标题',
                })
            continue  # 标题行不做后续检查

        # ---- 3) 引用格式检查 ----
        # 检测错误格式
        for bad_re, msg in _BAD_CITATION_PATTERNS:
            for m in bad_re.finditer(line):
                snippet = line[max(0, m.start() - 10):m.end() + 10]
                issues.append({
                    'level': 'error',
                    'category': '引用格式',
                    'message': f'第 {line_no} 行：{msg}：...{snippet}...',
                    'suggestion': '正文引用应使用英文方括号+英文逗号，如 [1,2] [3-5]',
                })

        # 检测合法引用中的编号顺序
        for m in _VALID_CITATION_RE.finditer(line):
            stats['citation_count'] += 1
            inner = m.group(1)
            nums = []
            for part in re.split(r'\s*,\s*', inner):
                range_match = re.match(r'(\d+)\s*[-–]\s*(\d+)', part)
                if range_match:
                    nums.extend(range(int(range_match.group(1)), int(range_match.group(2)) + 1))
                else:
                    nums.append(int(part.strip()))
            if len(nums) >= 2:
                for i in range(1, len(nums)):
                    if nums[i] < nums[i - 1]:
                        issues.append({
                            'level': 'warning',
                            'category': '引用格式',
                            'message': f'第 {line_no} 行：引用编号未按升序排列：[{inner}]',
                            'suggestion': '多引用应按编号升序排列，如 [3,5,7] 而非 [7,3,5]',
                        })
                        break

        # ---- 4) 写作风格检查（跳过引用块） ----
        if not line.lstrip().startswith('>'):
            for pattern, level, category, message, suggestion in _STYLE_CHECKS:
                for m in pattern.finditer(line):
                    start = max(0, m.start() - 8)
                    end = min(len(line), m.end() + 8)
                    snippet = line[start:end]
                    issues.append({
                        'level': level,
                        'category': category,
                        'message': f'第 {line_no} 行：{message}：...{snippet}...',
                        'suggestion': suggestion,
                    })

        # ---- 5) 图表编号收集 ----
        for m in fig_re.finditer(line):
            fig_refs.append((int(m.group(1)), int(m.group(2))))
        for m in tbl_re.finditer(line):
            tbl_refs.append((int(m.group(1)), int(m.group(2))))

        # A5 目标集合（正向）：行首题注无条件算定义
        m_cap = _A5_CAPTION_DEF_RE.match(line)
        if m_cap:
            num = m_cap.group(2).replace('–', '-')
            (fig_targets if m_cap.group(1) == '图' else tbl_targets).add(num)

        # A5 目标集合（减法）：剥离引用语境（含并列项）后，
        # 剩下的 图N-M / 表N-M 视为定义点（正文首述等）。
        line_defs = _A5_FT_REF_RE.sub('', line)
        for m in fig_re.finditer(line_defs):
            fig_targets.add(f"{m.group(1)}-{m.group(2)}")
        for m in tbl_re.finditer(line_defs):
            tbl_targets.add(f"{m.group(1)}-{m.group(2)}")

        # ---- 5b) A5 交叉引用收集（正文行）----
        for m in _A5_SEC_REF_RE.finditer(line):
            crossrefs.append((line_no, 'section', m.group(1), m.group(0)))
        for m in _A5_CHAP_REF_RE.finditer(line):
            crossrefs.append((line_no, 'section', m.group(1), m.group(0)))
        for m in _A5_EN_SEC_RE.finditer(line):
            crossrefs.append((line_no, 'section', m.group(1), m.group(0)))
        for m in _A5_FT_REF_RE.finditer(line):
            snippet = m.group(0)
            for im in _A5_FT_ITEM_RE.finditer(snippet):
                kind = 'figure' if im.group(1) == '图' else 'table'
                crossrefs.append((line_no, kind, im.group(2).replace('–', '-'), snippet))
        for m in _A5_EN_FIG_RE.finditer(line):
            crossrefs.append((line_no, 'figure', m.group(1).replace('–', '-'), m.group(0)))
        for m in _A5_EN_TBL_RE.finditer(line):
            crossrefs.append((line_no, 'table', m.group(1).replace('–', '-'), m.group(0)))
        for m in _A5_APP_REF_RE.finditer(line):
            crossrefs.append((line_no, 'appendix', m.group(1), m.group(0)))

        # ---- 6) 列表项检测 ----
        if list_re.match(line):
            issues.append({
                'level': 'warning',
                'category': '列表项',
                'message': f'第 {line_no} 行：检测到列表项格式',
                'suggestion': '学位论文应使用段落叙述，避免使用列表项',
            })

        # ---- 7) Pipe 表格语法检查 ----
        # 检测非标准表格写法（如 HTML <table> 或缩进式表格）
        if '<table' in line.lower() or '<tr' in line.lower():
            issues.append({
                'level': 'error',
                'category': '表格语法',
                'message': f'第 {line_no} 行：检测到 HTML 表格标签',
                'suggestion': '请使用 Markdown pipe 表格语法（| 列1 | 列2 |）',
            })

        # ---- 8) 图表题注格式检查 ----
        # 正确格式：「表 2-1：标题」或「图 3-2：标题」（中文冒号，章-序号）
        caption_line_re = re.match(r'^(图|表)\s*(\d+[-\u2013]\d+)\s*([：:])(.*)$', line)
        if caption_line_re:
            colon_char = caption_line_re.group(3)
            if colon_char == ':':
                issues.append({
                    'level': 'warning',
                    'category': '题注格式',
                    'message': f'第 {line_no} 行：题注使用了英文冒号',
                    'suggestion': '图表题注应使用中文冒号（：），如「表 2-1：实验结果」',
                })
            caption_text = caption_line_re.group(4).strip()
            if not caption_text:
                issues.append({
                    'level': 'error',
                    'category': '题注格式',
                    'message': f'第 {line_no} 行：题注缺少标题文字',
                    'suggestion': '图表题注冒号后应有描述性标题文字',
                })

        # ---- 9) [图]/[实验] 占位标记检查 ----
        placeholder_re = re.findall(r'\[(图[^]]*)\]|\[(实验[^]]*)\]', line)
        for match_groups in placeholder_re:
            marker = match_groups[0] or match_groups[1]
            issues.append({
                'level': 'info',
                'category': '占位标记',
                'message': f'第 {line_no} 行：检测到占位标记 [{marker}]',
                'suggestion': '请确认占位标记已替换为实际内容或正确的图表引用',
            })

    # ---- 2-续) 标题层级跳跃检查 ----
    for idx in range(1, len(heading_levels_seen)):
        prev_line, prev_lvl = heading_levels_seen[idx - 1]
        cur_line, cur_lvl = heading_levels_seen[idx]
        if cur_lvl > prev_lvl + 1:
            issues.append({
                'level': 'warning',
                'category': '标题层级',
                'message': f'第 {cur_line} 行：标题层级从 {prev_lvl} 级跳到 {cur_lvl} 级',
                'suggestion': '标题层级不应跳跃，如 # 后应先用 ## 再用 ###',
            })

    # ---- 5-续) 图表编号连续性检查 ----
    stats['figure_count'] = len(fig_refs)
    stats['table_count'] = len(tbl_refs)

    def _check_numbering(refs, label):
        """按章分组检查编号是否从 1 开始且连续"""
        by_chapter = {}
        for chap, seq in refs:
            by_chapter.setdefault(chap, []).append(seq)
        for chap, seqs in sorted(by_chapter.items()):
            seqs_sorted = sorted(seqs)
            expected = list(range(1, len(seqs_sorted) + 1))
            if seqs_sorted != expected:
                issues.append({
                    'level': 'warning',
                    'category': '图表编号',
                    'message': f'第 {chap} 章{label}编号不连续：实际为 {seqs_sorted}，期望为 {expected}',
                    'suggestion': f'{label}编号应从 1 开始且连续递增',
                })

    _check_numbering(fig_refs, '图')
    _check_numbering(tbl_refs, '表')

    # ---- A5) 章节交叉引用有效性（断链 → WARN，保守）----
    # 引用目标在全文找不到对应定义（章节标题 / 图表编号 / 附录）→ 断链。
    # 提取易误判（如把正文里偶发的"第3节"当引用），故只 WARN，不阻断。
    _a5_target_map = {
        'section': (section_targets, '节'),
        'figure': (fig_targets, '图'),
        'table': (tbl_targets, '表'),
        'appendix': (appendix_targets, '附录'),
    }
    for ref_line, kind, target, snippet in crossrefs:
        targets, label = _a5_target_map[kind]
        if target not in targets:
            issues.append({
                'level': 'warning',
                'category': '交叉引用',
                # code 供 Step 9 以 issue_summary.xref_broken 做 HALT 判据。
                # level 恒为 warning、不进硬阻断 code 集合：退出码语义零变化。
                'code': 'xref_broken',
                'message': f'第 {ref_line} 行：交叉引用 "{snippet}" 指向的{label} {target} 在全文中找不到定义',
                'suggestion': f'核对该{label}是否存在；断链请修正编号或补全被引用的{label}',
            })

    # ---- 10) 题名长度检查（≤25 字） ----
    for line_no, raw_line in enumerate(lines, start=1):
        h1_match = re.match(r'^#\s+(.+)', raw_line.rstrip('\n'))
        if h1_match:
            title_text = h1_match.group(1).strip()
            # 只计中文字符 + 英文单词（每个单词算 1 字）
            cn_count = len(re.findall(r'[\u4e00-\u9fff]', title_text))
            if cn_count > 25:
                issues.append({
                    'level': 'warning',
                    'category': '题名长度',
                    'message': f'第 {line_no} 行：一级标题中文字数为 {cn_count}（超过 25 字上限）',
                    'suggestion': '建议论文题名一般不超过 25 个汉字',
                })
            break  # 只检查首个 H1

    # ---- 11) 关键词分隔符检查 ----
    for line_no, raw_line in enumerate(lines, start=1):
        line = raw_line.rstrip('\n')
        # 中文关键词行
        cn_kw_match = re.match(r'^关键词[：:]\s*(.+)', line)
        if cn_kw_match:
            kw_text = cn_kw_match.group(1)
            if ';' in kw_text or ',' in kw_text or '、' in kw_text:
                issues.append({
                    'level': 'warning',
                    'category': '关键词格式',
                    'message': f'第 {line_no} 行：中文关键词应使用全角分号（；）分隔',
                    'suggestion': '中文关键词之间用全角分号"；"分隔，如"机器学习；深度学习；神经网络"',
                })
        # 英文关键词行
        en_kw_match = re.match(r'^[Kk]eywords?[：:]\s*(.+)', line)
        if en_kw_match:
            kw_text = en_kw_match.group(1)
            if '；' in kw_text or '，' in kw_text or '、' in kw_text:
                issues.append({
                    'level': 'warning',
                    'category': '关键词格式',
                    'message': f'第 {line_no} 行：英文关键词应使用半角分号（;）分隔',
                    'suggestion': '英文关键词之间用半角分号";"分隔，如"machine learning; deep learning"',
                })

    if md_checks == 'xref':
        issues = [i for i in issues if i.get('category') == '交叉引用']

    return issues, stats


def check_caption_style(doc, format_context=None):
    """
    检查图表题注格式是否符合规范：
    - 字体：楷体 (KaiTi)
    - 字号：10.5pt（五号）
    - 对齐：居中
    - 行距：单倍行距
    - 段后间距：12pt
    """
    issues = []
    reported = set()
    FONT_SIZE_TOL = 0.5
    SPACING_TOL = 1.0
    format_context = format_context or build_format_render_context({})

    caption_re = re.compile(r'^(图|表)\s*\d')

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        if not caption_re.match(text):
            continue

        cap_type = '图题注' if text.startswith('图') else '表题注'
        spec_key = "figure_caption" if text.startswith("图") else "table_caption"
        spec = format_context.get("style_profile", {}).get(spec_key, {})

        # 字号检查
        if para.runs:
            run = para.runs[0]
            if run.font.size is not None and hasattr(run.font.size, 'pt'):
                actual = run.font.size.pt
                expected_size = spec.get("font_size_pt", 10.5)
                if abs(actual - expected_size) > FONT_SIZE_TOL:
                    tag = f'caption_font_size'
                    if tag not in reported:
                        reported.add(tag)
                        issues.append({
                            'level': 'error', 'category': '题注格式',
                            'message': f'{cap_type}字号不符：{actual}pt（应为 {expected_size}pt）',
                            'suggestion': f'图表题注字号应设置为 {expected_size}pt'
                        })

            # 字体检查
            ea_font = _get_east_asian_font(run)
            if ea_font is not None:
                tag = f'caption_font_ea'
                expected_ea = spec.get("font_east_asia", "KaiTi")
                if tag not in reported and ea_font != expected_ea:
                    reported.add(tag)
                    issues.append({
                        'level': 'error', 'category': '题注格式',
                        'message': f'{cap_type}中文字体不符：{ea_font}（应为 {expected_ea}）',
                        'suggestion': f'图表题注中文字体应设置为 {expected_ea}'
                    })

        # 对齐检查（居中）
        actual_align = para.paragraph_format.alignment
        if actual_align is not None:
            align_val = actual_align if isinstance(actual_align, int) else actual_align.value if hasattr(actual_align, 'value') else None
            expected_align = _align_value(spec.get("alignment", "center"))
            if align_val is not None and align_val != expected_align:
                tag = f'caption_alignment'
                if tag not in reported:
                    reported.add(tag)
                    issues.append({
                        'level': 'error', 'category': '题注格式',
                        'message': f'{cap_type}对齐方式不符',
                        'suggestion': f'图表题注应设置为 {spec.get("alignment", "center")} 对齐'
                    })

        # 段前/段后间距检查（图题注：段前0/段后12pt；表题注：段前12pt/段后0）
        is_figure = text.startswith('图')
        sb = para.paragraph_format.space_before
        sa = para.paragraph_format.space_after
        if is_figure:
            # 图题注：段前0，段后12pt（段后1行）
            expected_after = spec.get("space_after_pt", 12.0)
            if sa is not None and hasattr(sa, 'pt'):
                actual_sa = sa.pt
                if abs(actual_sa - expected_after) > SPACING_TOL:
                    tag = f'caption_fig_space_after'
                    if tag not in reported:
                        reported.add(tag)
                        issues.append({
                            'level': 'warning', 'category': '题注格式',
                            'message': f'图题注段后间距不符：{actual_sa:.1f}pt（应为 {expected_after}pt）',
                            'suggestion': f'图题注段后间距应设置为 {expected_after}pt'
                        })
            expected_before = spec.get("space_before_pt", 0.0)
            if sb is not None and hasattr(sb, 'pt'):
                actual_sb = sb.pt
                if abs(actual_sb - expected_before) > SPACING_TOL:
                    tag = f'caption_fig_space_before'
                    if tag not in reported:
                        reported.add(tag)
                        issues.append({
                            'level': 'warning', 'category': '题注格式',
                            'message': f'图题注段前间距不符：{actual_sb:.1f}pt（应为 {expected_before}pt）',
                            'suggestion': f'图题注段前间距应设置为 {expected_before}pt'
                        })
        else:
            # 表题注：段前12pt（段前1行），段后0
            expected_before = spec.get("space_before_pt", 12.0)
            if sb is not None and hasattr(sb, 'pt'):
                actual_sb = sb.pt
                if abs(actual_sb - expected_before) > SPACING_TOL:
                    tag = f'caption_tbl_space_before'
                    if tag not in reported:
                        reported.add(tag)
                        issues.append({
                            'level': 'warning', 'category': '题注格式',
                            'message': f'表题注段前间距不符：{actual_sb:.1f}pt（应为 {expected_before}pt）',
                            'suggestion': f'表题注段前间距应设置为 {expected_before}pt'
                        })
            expected_after = spec.get("space_after_pt", 0.0)
            if sa is not None and hasattr(sa, 'pt'):
                actual_sa = sa.pt
                if abs(actual_sa - expected_after) > SPACING_TOL:
                    tag = f'caption_tbl_space_after'
                    if tag not in reported:
                        reported.add(tag)
                        issues.append({
                            'level': 'warning', 'category': '题注格式',
                            'message': f'表题注段后间距不符：{actual_sa:.1f}pt（应为 {expected_after}pt）',
                            'suggestion': f'表题注段后间距应设置为 {expected_after}pt'
                        })

        # 行距检查（单倍行距 ≈ 12pt for 10.5pt font, or linespacing rule SINGLE）
        pf = para.paragraph_format
        if pf.line_spacing_rule is not None:
            rule_val = pf.line_spacing_rule if isinstance(pf.line_spacing_rule, int) else pf.line_spacing_rule.value if hasattr(pf.line_spacing_rule, 'value') else None
            # WD_LINE_SPACING.SINGLE = 0, EXACTLY = 4, AT_LEAST = 3, MULTIPLE = 5
            if rule_val is not None and rule_val not in (0, None):
                # 如果不是单倍行距，检查是否为固定值且接近单倍
                actual_ls = line_spacing_pt(para)
                if actual_ls is not None and actual_ls > 15.0:
                    tag = f'caption_line_spacing'
                    if tag not in reported:
                        reported.add(tag)
                        issues.append({
                            'level': 'warning', 'category': '题注格式',
                            'message': f'{cap_type}行距不符：{actual_ls:.1f}pt（应为单倍行距）',
                            'suggestion': '图表题注行距应设置为单倍行距'
                        })

    return issues


def check_table_cell_format(doc, format_context=None):
    """
    检查表格单元格格式：
    - 字体：宋体 (SimSun) 10.5pt（五号）
    - 表头行加粗
    - 单元格居中对齐
    """
    issues = []
    reported = set()
    FONT_SIZE_TOL = 0.5
    format_context = format_context or build_format_render_context({})
    cell_spec = format_context.get("style_profile", {}).get("table_cell", {})
    expected_align = _align_value(cell_spec.get("alignment", "center"))

    for t_idx, table in enumerate(doc.tables):
        num_rows = len(table.rows)
        if num_rows < 2:
            continue

        table_label = f'表格 {t_idx + 1}'

        for r_idx, row in enumerate(table.rows):
            is_header = (r_idx == 0)
            for c_idx, cell in enumerate(row.cells):
                for para in cell.paragraphs:
                    if not para.text.strip():
                        continue

                    # 居中对齐检查
                    actual_align = para.paragraph_format.alignment
                    if actual_align is not None:
                        align_val = actual_align if isinstance(actual_align, int) else actual_align.value if hasattr(actual_align, 'value') else None
                        if align_val is not None and align_val != expected_align:
                            tag = f'tbl_cell_align_{t_idx}'
                            if tag not in reported:
                                reported.add(tag)
                                issues.append({
                                    'level': 'warning', 'category': '表格单元格',
                                    'message': f'{table_label}：单元格对齐方式不符',
                                    'suggestion': f'三线表单元格内容应设置为 {cell_spec.get("alignment", "center")} 对齐'
                                })

                    for run in para.runs:
                        # 字号检查
                        if run.font.size is not None and hasattr(run.font.size, 'pt'):
                            actual = run.font.size.pt
                            expected_size = cell_spec.get("font_size_pt", 10.5)
                            if abs(actual - expected_size) > FONT_SIZE_TOL:
                                tag = f'tbl_cell_fontsize_{t_idx}'
                                if tag not in reported:
                                    reported.add(tag)
                                    issues.append({
                                        'level': 'warning', 'category': '表格单元格',
                                        'message': f'{table_label}：单元格字号不符 {actual}pt（应为 {expected_size}pt）',
                                        'suggestion': f'表格单元格字号应设置为 {expected_size}pt'
                                    })

                        # 中文字体检查
                        ea_font = _get_east_asian_font(run)
                        expected_ea = cell_spec.get("font_east_asia", "SimSun")
                        if ea_font is not None and ea_font != expected_ea:
                            tag = f'tbl_cell_fontea_{t_idx}'
                            if tag not in reported:
                                reported.add(tag)
                                issues.append({
                                    'level': 'warning', 'category': '表格单元格',
                                    'message': f'{table_label}：单元格中文字体不符 {ea_font}（应为 {expected_ea}）',
                                    'suggestion': f'表格单元格中文字体应设置为 {expected_ea}'
                                })

                        # 表头加粗检查
                        if is_header:
                            expected_header_bold = cell_spec.get("header_bold", True)
                            if run.font.bold is not None and run.font.bold != expected_header_bold:
                                tag = f'tbl_header_bold_{t_idx}'
                                if tag not in reported:
                                    reported.add(tag)
                                    issues.append({
                                        'level': 'warning', 'category': '表格单元格',
                                        'message': f'{table_label}：表头行加粗属性不符',
                                        'suggestion': f'三线表表头行文字应设置为 {"加粗" if expected_header_bold else "常规"}'
                                    })
                        break  # 只检查第一个 run
                    break  # 只检查第一个段落

    return issues


def check_run_level_font_pairing(doc):
    """
    检查正文 run 级别的双字体配对完整性：
    每个 run 应同时设置东亚字体（宋体）和西文字体（Times New Roman）。
    仅抽样检查前 20 个正文段落。
    """
    issues = []
    sample_count = 0
    max_samples = 20
    missing_ea = 0
    missing_latin = 0

    for para in doc.paragraphs:
        if not para.text.strip():
            continue
        style_name = para.style.name if para.style else None
        if style_name not in ('Normal', '正文'):
            continue

        sample_count += 1
        if sample_count > max_samples:
            break

        for run in para.runs:
            if not run.text.strip():
                continue
            # 检查西文字体
            if run.font.name is None:
                missing_latin += 1
            # 检查东亚字体
            ea_font = _get_east_asian_font(run)
            if ea_font is None:
                missing_ea += 1

    if missing_ea > 3:
        issues.append({
            'level': 'warning', 'category': '字体配对',
            'message': f'正文中 {missing_ea} 个 run 缺少东亚字体定义（抽样 {max_samples} 段）',
            'suggestion': '每个 run 应同时设置东亚字体（宋体）和西文字体（Times New Roman）'
        })
    if missing_latin > 3:
        issues.append({
            'level': 'warning', 'category': '字体配对',
            'message': f'正文中 {missing_latin} 个 run 缺少西文字体定义（抽样 {max_samples} 段）',
            'suggestion': '每个 run 应同时设置东亚字体（宋体）和西文字体（Times New Roman）'
        })

    return issues


def check_header_footer(doc, format_context=None):
    """
    检查页眉页脚格式是否符合当前配置：
    - 页眉：宋体五号(10.5pt)，左侧为配置中的页眉文案，距顶端按配置
    - 页脚：TNR 小五号(9pt)，居中页码，距底端1.75cm
    """
    issues = []
    FONT_SIZE_TOL = 0.5
    DISTANCE_TOL = Cm(0.2)
    format_context = format_context or {}
    _degree_label = format_context.get("degree_type", "博士学位论文")
    _univ = format_context.get("university_name", "示例大学")
    _default_header = f"{_univ}{_degree_label}"
    expected_header_left_text = format_context.get("header_left_text", _default_header)
    expected_header_distance_cm = format_context.get("header_distance_cm", 1.5)
    expected_footer_distance_cm = format_context.get("footer_distance_cm", 1.75)
    style_profile = format_context.get("style_profile", {}) if isinstance(format_context, dict) else {}
    header_spec = style_profile.get("header", {})
    footer_spec = style_profile.get("footer", {})

    for sec_idx, section in enumerate(doc.sections):
        sec_label = f'第 {sec_idx + 1} 节'

        # ---- 页眉检查 ----
        header = section.header
        if header and not header.is_linked_to_previous:
            h_paras = [p for p in header.paragraphs if p.text.strip()]
            if not h_paras:
                issues.append({
                    'level': 'warning', 'category': '页眉格式',
                    'message': f'{sec_label}：页眉为空',
                    'suggestion': f'页眉应包含"{expected_header_left_text}"（左）和章名（右）',
                })
            else:
                h_text = h_paras[0].text.strip()
                if expected_header_left_text not in h_text:
                    issues.append({
                        'level': 'warning', 'category': '页眉格式',
                        'message': f'{sec_label}：页眉缺少"{expected_header_left_text}"',
                        'suggestion': f'页眉左侧应为"{expected_header_left_text}"',
                    })
                # 字号检查
                for run in h_paras[0].runs:
                    if run.font.size is not None:
                        actual_pt = run.font.size.pt
                        expected_header_size = header_spec.get("font_size_pt", 10.5)
                        if abs(actual_pt - expected_header_size) > FONT_SIZE_TOL:
                            issues.append({
                                'level': 'warning', 'category': '页眉格式',
                                'message': f'{sec_label}：页眉字号为 {actual_pt}pt，应为 {expected_header_size}pt',
                                'suggestion': f'页眉字号应为 {expected_header_size}pt',
                            })
                        break

        # 页眉距顶端
        if section.header_distance is not None:
            expected = Cm(expected_header_distance_cm)
            if abs(section.header_distance - expected) > DISTANCE_TOL:
                actual_cm = section.header_distance / Cm(1)
                issues.append({
                    'level': 'warning', 'category': '页眉格式',
                    'message': f'{sec_label}：页眉距顶端 {actual_cm:.2f}cm，应为 {expected_header_distance_cm}cm',
                    'suggestion': f'页眉距顶端应设置为 {expected_header_distance_cm}cm',
                })

        # ---- 页脚检查 ----
        footer = section.footer
        if footer and not footer.is_linked_to_previous:
            f_paras = [p for p in footer.paragraphs if p.text.strip() or
                       any(el.tag.endswith('fldChar') for el in p._element.iter())]
            if not f_paras:
                issues.append({
                    'level': 'warning', 'category': '页脚格式',
                    'message': f'{sec_label}：页脚为空（缺少页码）',
                    'suggestion': '页脚应包含居中页码',
                })
            else:
                # 居中检查
                p = f_paras[0]
                if p.alignment not in (WD_ALIGN_PARAGRAPH.CENTER, None):
                    issues.append({
                        'level': 'warning', 'category': '页脚格式',
                        'message': f'{sec_label}：页脚未居中对齐',
                        'suggestion': '页码应居中对齐',
                    })
                # 字号检查
                for run in p.runs:
                    if run.font.size is not None:
                        actual_pt = run.font.size.pt
                        expected_footer_size = footer_spec.get("font_size_pt", 9)
                        if abs(actual_pt - expected_footer_size) > FONT_SIZE_TOL:
                            issues.append({
                                'level': 'warning', 'category': '页脚格式',
                                'message': f'{sec_label}：页脚字号为 {actual_pt}pt，应为 {expected_footer_size}pt',
                                'suggestion': f'页码字号应为 {expected_footer_size}pt',
                            })
                        break

        # 页脚距底端
        if section.footer_distance is not None:
            expected = Cm(expected_footer_distance_cm)
            if abs(section.footer_distance - expected) > DISTANCE_TOL:
                actual_cm = section.footer_distance / Cm(1)
                issues.append({
                    'level': 'warning', 'category': '页脚格式',
                    'message': f'{sec_label}：页脚距底端 {actual_cm:.2f}cm，应为 {expected_footer_distance_cm}cm',
                    'suggestion': f'页脚距底端应设置为 {expected_footer_distance_cm}cm',
                })

    return issues


def generate_quality_report(
    docx_path,
    verbose=True,
    body_target_chars=80000,
    review_target_chars=0,
    review_in_scope=False,
    references_min_count=80,
    min_chapters=5,
    enforce_full_structure=False,
    md_path=None,
    per_chapter_ref_floor=None,
    md_checks='all',
):
    """生成完整质量报告"""
    gate_payload = _pending_template_payload(docx_path)
    if gate_payload is not None:
        return gate_payload
    try:
        doc = Document(docx_path)
    except Exception as e:
        return {
            'success': False,
            'error': f'无法打开文件：{str(e)}'
        }
    format_context = _load_render_context_for_docx(docx_path)
    
    all_issues = []
    
    # 1. 检查字数
    if verbose:
        print("🔍 检查字数...")
    word_issues, word_stats = check_word_count(
        doc,
        body_target_chars=body_target_chars,
        review_target_chars=review_target_chars,
        review_in_scope=review_in_scope,
    )
    all_issues.extend(word_issues)
    
    # 2. 检查标题层级
    if verbose:
        print("🔍 检查标题层级...")
    heading_issues = check_heading_levels(doc)
    all_issues.extend(heading_issues)
    
    # 3. 检查图表编号
    if verbose:
        print("🔍 检查图表编号...")
    numbering_issues = check_figure_numbering(doc)
    all_issues.extend(numbering_issues)
    
    # 4. 检查列表项
    if verbose:
        print("🔍 检查列表项...")
    bullet_issues = check_bullet_points(doc)
    all_issues.extend(bullet_issues)
    
    # 5. 检查参考文献
    if verbose:
        print("🔍 检查参考文献...")
    ref_issues, ref_count = check_reference_count(doc, min_reference_count=references_min_count)
    all_issues.extend(ref_issues)

    # 5.0.1 按章软文献分布（warning，不阻断）
    all_issues.extend(check_chapter_reference_distribution(doc, per_chapter_ref_floor))

    # 5.1 参考文献位置校验
    reference_position_issues = check_reference_position(doc)
    all_issues.extend(reference_position_issues)

    # 5.1.1 参考文献著录格式校验（从 literature_index.json 读取）
    _lit_index_path = None
    _inferred_root = infer_project_root_for_profile(docx_path)
    if _inferred_root:
        _candidate = os.path.join(_inferred_root, "literature_index.json")
        if os.path.exists(_candidate):
            _lit_index_path = _candidate
    if _lit_index_path and _validate_references is not None:
        if verbose:
            print("🔍 检查参考文献著录格式...")
        try:
            with open(_lit_index_path, "r", encoding="utf-8") as _f:
                _lit_entries = json.load(_f)
            if isinstance(_lit_entries, list):
                _ref_format_issues = _validate_references(_lit_entries)
                # 映射字段到 check_quality 统一格式
                for _iss in _ref_format_issues:
                    all_issues.append({
                        "level": _iss.get("level", "warning"),
                        "category": "参考文献著录",
                        "message": _iss.get("message", ""),
                        "suggestion": "补齐 literature_index.json 中缺失字段，或修正 type 标识；详见 scripts/reference_renderer.py",
                    })
        except Exception as _e:
            all_issues.append({
                "level": "warning",
                "category": "参考文献著录",
                "message": f"无法读取 literature_index.json 进行著录格式校验：{_e}",
                "suggestion": "确认 literature_index.json 存在且为合法 JSON 数组",
            })

    # 5.2 正文内联引用格式检测
    if verbose:
        print("🔍 检查正文引用格式...")
    citation_issues = check_inline_citations(doc)
    all_issues.extend(citation_issues)

    # 5.3 写作风格检测
    if verbose:
        print("🔍 检查写作风格...")
    style_issues = check_writing_style(doc)
    all_issues.extend(style_issues)

    # 5.3.1 字符级检查（D1 半角标点 / D2 上下标裸写 / F1 中文错别字，全 WARN）
    if verbose:
        print("🔍 字符级检查（半角标点/上下标/错别字）...")
    char_issues = check_char_level(doc)
    all_issues.extend(char_issues)

    # 5.4 全文结构门禁（仅在全文检查时启用）
    if enforce_full_structure:
        if verbose:
            print("🔍 检查全文结构门禁...")
        structure_issues = check_full_thesis_structure(doc, min_chapters=min_chapters)
        all_issues.extend(structure_issues)

    # 5.5 全文结构顺序校验（仅在全文检查时启用）
    if enforce_full_structure:
        if verbose:
            print("🔍 检查全文结构顺序...")
        order_issues = check_section_order(doc)
        all_issues.extend(order_issues)

    # 5.6 章间分页符校验（仅在全文检查时启用）
    if enforce_full_structure:
        if verbose:
            print("🔍 检查章间分页符...")
        page_break_issues = check_page_breaks_between_chapters(doc)
        all_issues.extend(page_break_issues)

    # 5.7 章间衔接检查（防订书机论文，仅在全文检查时启用）
    if enforce_full_structure:
        if verbose:
            print("🔍 检查章间衔接...")
        coherence_issues = check_cross_chapter_coherence(
            doc, project_root=infer_project_root_for_profile(docx_path)
        )
        all_issues.extend(coherence_issues)
    
    # 6. 检查段落格式
    if verbose:
        print("🔍 检查段落格式...")
    format_issues = check_paragraph_formatting(doc)
    all_issues.extend(format_issues)

    # 7. 检查三线表格式
    if verbose:
        print("🔍 检查三线表格式...")
    table_issues = check_table_format(doc)
    all_issues.extend(table_issues)

    # 7.1 检查图编号映射一致性（如果 figure_map.json 存在）
    if verbose:
        print("🔍 检查图编号映射...")
    figure_map_project_root = infer_project_root_for_profile(docx_path)
    figure_map_issues = check_figure_map_consistency(doc, project_root=figure_map_project_root)
    all_issues.extend(figure_map_issues)

    # 8. 检查缩略语一致性
    if verbose:
        print("🔍 检查缩略语一致性...")
    abbr_project_root = infer_project_root_for_profile(docx_path)
    abbr_issues = check_abbreviation_consistency(doc, project_root=abbr_project_root)
    all_issues.extend(abbr_issues)

    # 9. Word 格式合规检查（页面布局、字体、字号、行距等）
    if verbose:
        print("🔍 检查 Word 格式合规...")
    word_format_issues = check_word_format_compliance(doc, format_context=format_context)
    all_issues.extend(word_format_issues)

    # 9.1 题注格式检查（楷体/10.5pt/居中/单倍行距/段后12pt）
    if verbose:
        print("🔍 检查题注格式...")
    caption_issues = check_caption_style(doc, format_context=format_context)
    all_issues.extend(caption_issues)

    # 9.2 表格单元格格式检查（宋体10.5pt/表头加粗/居中）
    if verbose:
        print("🔍 检查表格单元格格式...")
    cell_issues = check_table_cell_format(doc, format_context=format_context)
    all_issues.extend(cell_issues)

    # 9.3 Run 级别双字体配对完整性检查
    if verbose:
        print("🔍 检查字体配对完整性...")
    pairing_issues = check_run_level_font_pairing(doc)
    all_issues.extend(pairing_issues)

    # 9.4 页眉页脚格式检查
    if verbose:
        print("🔍 检查页眉页脚格式...")
    hf_issues = check_header_footer(doc, format_context=format_context)
    all_issues.extend(hf_issues)

    # 10. Markdown 质量检查（如果提供了 md_path）
    md_stats = None
    if md_path:
        if verbose:
            print("🔍 检查 Markdown 质量...")
        md_issues, md_stats = check_markdown_quality(md_path, md_checks=md_checks)
        all_issues.extend(md_issues)
    
    # 计算总分（简化评分）
    error_count = len([i for i in all_issues if i['level'] == 'error'])
    # 按章文献密度为软提醒(soft target)，不计入导出评分，避免累积软告警把总分压到 80 门槛下、
    # 间接把"软目标"变成硬拦(与用户"软目标只提醒不阻断"的决策一致)。
    warning_count = len([i for i in all_issues
                         if i['level'] == 'warning' and i.get('category') != '文献引用分布'])
    info_count = len([i for i in all_issues if i['level'] == 'info'])
    # 上下标裸写 / 中文句内半角标点 为硬阻断项：单独计数，纳入 ok/退出码判定（不改其评分权重）
    subsup_bare_count = len([i for i in all_issues if i.get('code') == 'subsup_bare'])
    halfwidth_in_cn_count = len([i for i in all_issues if i.get('code') == 'halfwidth_punct_in_cn'])
    english_misspelling_count = len([i for i in all_issues if i.get('code') == 'english_misspelling'])
    # 破折号(——)硬门禁、禁止使用：命中即非零退出，无论总分高低（与上下标/半角同列硬阻断）
    em_dash_count = len([i for i in all_issues
                         if i.get('category') == '标点规范' and i.get('message', '').startswith('使用了破折号')])
    # 去AI必禁三项之另两项：scare quotes（引号规范）+ 解释性冒号（标点规范），与破折号同级硬阻断
    scare_quote_count = len([i for i in all_issues
                             if i.get('category') == '引号规范' and i.get('message', '').startswith('疑似 scare quotes')])
    explanatory_colon_count = len([i for i in all_issues
                                   if i.get('category') == '标点规范' and i.get('message', '').startswith('疑似解释性冒号')])
    # A5 交叉引用断链：单列计数供 Step 9 HALT 判据用。**不是硬阻断项**——不进退出码判定，
    # 一条断链只按 warning 扣 3 分（A5 是启发式提取，硬拦一次假阳就卡死全文总检）。
    xref_broken_count = len([i for i in all_issues if i.get('code') == 'xref_broken'])

    total_score = 100 - error_count * 10 - warning_count * 3 - info_count * 1
    total_score = max(0, min(100, total_score))
    
    report = {
        'success': True,
        'file_path': docx_path,
        'check_date': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
        'overall_score': total_score,
        'statistics': {
            'body_words': word_stats['body_words'],
            'review_words': word_stats['review_words'],
            'reference_count': ref_count,
            'reference_target_count': int(references_min_count),
            'total_paragraphs': len(doc.paragraphs),
            'total_tables': len(doc.tables),
            'markdown': md_stats,
        },
        'issue_summary': {
            'total': len(all_issues),
            'error': error_count,
            'warning': warning_count,
            'info': info_count,
            'subsup_bare': subsup_bare_count,
            'halfwidth_punct_in_cn': halfwidth_in_cn_count,
            'english_misspelling': english_misspelling_count,
            'decorative_em_dash': em_dash_count,
            'scare_quotes': scare_quote_count,
            'explanatory_colon': explanatory_colon_count,
            'xref_broken': xref_broken_count
        },
        'issues': all_issues,
        'targets': {
            'body_target_chars': int(body_target_chars),
            'review_target_chars': int(review_target_chars),
            'review_in_scope': bool(review_in_scope),
            'references_min_count': int(references_min_count),
            'min_chapters': int(min_chapters),
            'enforce_full_structure': bool(enforce_full_structure),
        },
        'recommendations': generate_recommendations(all_issues, word_stats)
    }
    
    return report


def generate_recommendations(issues, word_stats):
    """根据检查结果生成建议"""
    recommendations = []
    
    # 字数建议
    body_target = int(word_stats.get('body_target_chars', 80000) or 80000)
    if word_stats['body_words'] < body_target:
        gap = body_target - word_stats['body_words']
        recommendations.append(
            f"正文字数不足 {gap} 字，建议扩展以下内容：\n"
            "  - 增加机制讨论（从分子、细胞、体内三个层面）\n"
            "  - 补充文献对比分析\n"
            "  - 增加局限性讨论和未来展望"
        )
    
    # 错误处理建议
    errors = [i for i in issues if i['level'] == 'error']
    if errors:
        recommendations.append(
            f"发现 {len(errors)} 个严重错误，必须修正：\n" +
            "\n".join(f"  - {e['message']}" for e in errors[:3])
        )

    reference_errors = [i for i in errors if i.get('category') == '参考文献']
    if reference_errors:
        recommendations.append(
            "参考文献未达配置下限，建议补充近5年高质量文献并统一格式（GB/T 7714 或学校模板）。"
        )
    
    # 格式建议
    format_issues = [i for i in issues if i['category'] == '格式']
    if len(format_issues) > 5:
        recommendations.append(
            "格式问题较多，建议：\n"
            "  - 使用样式模板统一格式\n"
            "  - 检查所有段落的首行缩进和行距"
        )
    
    if not recommendations:
        recommendations.append("论文质量良好，符合基本要求！")
    
    return recommendations


def format_report_text(report):
    """格式化报告为可读文本"""
    lines = []
    lines.append("=" * 70)
    lines.append("📋 博士论文质量检查报告")
    lines.append("=" * 70)
    lines.append(f"📄 文件：{os.path.basename(report['file_path'])}")
    lines.append(f"📅 检查时间：{report['check_date']}")
    lines.append(f"⭐ 总体评分：{report['overall_score']} / 100")
    lines.append("")
    
    # 统计信息
    stats = report['statistics']
    targets = report.get('targets', {})
    body_target = int(targets.get('body_target_chars', stats.get('body_target_chars', 80000)) or 80000)
    review_target = int(targets.get('review_target_chars', stats.get('review_target_chars', 0)) or 0)
    review_in_scope = bool(targets.get('review_in_scope', stats.get('review_in_scope', False)))
    references_min_count = int(
        targets.get('references_min_count', stats.get('reference_target_count', 80)) or 80
    )
    min_chapters = int(targets.get('min_chapters', 5) or 5)
    enforce_full_structure = bool(targets.get('enforce_full_structure', False))
    lines.append("📊 统计信息：")
    lines.append(f"   正文字数：{stats['body_words']:,} / {body_target:,}")
    if review_in_scope and review_target > 0:
        lines.append(f"   综述字数：{stats['review_words']:,} / {review_target:,}")
    else:
        lines.append(f"   综述字数：{stats['review_words']:,}（不纳入当前考核）")
    if references_min_count > 0:
        lines.append(f"   参考文献：{stats['reference_count']} / {references_min_count}")
    else:
        lines.append(f"   参考文献：{stats['reference_count']}（当前未纳入阈值检查）")
    if enforce_full_structure:
        lines.append(f"   全文结构门禁：启用（最少章节 {min_chapters}）")
    else:
        lines.append("   全文结构门禁：未启用")
    lines.append(f"   总段落数：{stats['total_paragraphs']:,}")
    lines.append(f"   总表格数：{stats['total_tables']}")
    lines.append("")
    
    # 问题汇总
    summary = report['issue_summary']
    lines.append("⚠️  问题汇总：")
    lines.append(f"   严重错误：{summary['error']} 个")
    lines.append(f"   警告：{summary['warning']} 个")
    lines.append(f"   提示：{summary['info']} 个")
    lines.append("")
    
    # 详细问题列表
    if report['issues']:
        lines.append("📝 详细问题：")
        for i, issue in enumerate(report['issues'][:20], 1):  # 只显示前 20 个
            level_icon = {
                'error': '❌',
                'warning': '⚠️ ',
                'info': 'ℹ️ '
            }.get(issue['level'], '•')
            
            location = issue.get('location', '')
            location_str = f" [{location}]" if location else ""
            
            lines.append(f"{i}. {level_icon} {issue['message']}{location_str}")
            lines.append(f"   💡 {issue['suggestion']}")
            lines.append("")
        
        if len(report['issues']) > 20:
            lines.append(f"   ... 还有 {len(report['issues']) - 20} 个问题未显示")
            lines.append("")
    
    # 建议
    lines.append("💡 改进建议：")
    for i, rec in enumerate(report['recommendations'], 1):
        lines.append(f"{i}. {rec}")
        lines.append("")
    
    lines.append("=" * 70)
    
    return "\n".join(lines)


def main():
    """命令行入口"""
    parser = argparse.ArgumentParser(description="博士论文质量自检工具")
    parser.add_argument("docx_path", help="待检查的 docx 文件路径")
    parser.add_argument("--output", choices=["json", "text"], default="text", help="输出格式")
    parser.add_argument("--profile", help="thesis_profile.json 路径（可选）")
    parser.add_argument("--body-target", type=int, help="覆盖正文目标字数")
    parser.add_argument("--review-target", type=int, help="覆盖综述目标字数")
    parser.add_argument("--review-in-scope", action="store_true", help="将综述纳入考核目标")
    parser.add_argument("--references-min", type=int, help="覆盖最少参考文献数量")
    parser.add_argument("--min-chapters", type=int, help="覆盖最少章节数（全文结构门禁）")
    parser.add_argument("--enforce-full-structure", action="store_true", help="启用全文结构门禁（章数/首章绪论/末章总结）")
    parser.add_argument("--md", dest="md_path", help="同时检查对应的 Markdown 源文件")
    parser.add_argument("--md-checks", choices=["all", "xref"], default="all",
                        help="Markdown 侧检查范围：all=全部 13 类（默认，语义同历史）；"
                             "xref=只报交叉引用断链（全文总检用的窄口）")
    args = parser.parse_args()
    docx_path = args.docx_path
    output_format = args.output
    
    if not os.path.exists(docx_path):
        if output_format == 'json':
            print(json.dumps({
                "success": False,
                "error": "file_not_found",
                "file_path": docx_path,
                "message": f"文件不存在：{docx_path}",
            }, ensure_ascii=False, indent=2))
        else:
            print(f"❌ 文件不存在：{docx_path}")
        sys.exit(1)
    
    verbose = output_format != 'json'
    if verbose:
        print("🔍 开始质量检查...\n")

    profile_project_root = infer_project_root_for_profile(docx_path)
    try:
        profile, _ = load_profile(profile_project_root, args.profile)
    except Exception as e:
        payload = {
            "success": False,
            "error": "profile_load_failed",
            "message": str(e),
            "profile_path": args.profile,
        }
        if output_format == 'json':
            print(json.dumps(payload, ensure_ascii=False, indent=2))
        else:
            print(f"❌ 配置加载失败：{payload['message']}")
        sys.exit(1)
    targets = profile.get("targets", {}) if isinstance(profile, dict) else {}
    _fmt_profile = profile.get("format_profile", {}) if isinstance(profile, dict) else {}
    _degree_type = str(_fmt_profile.get("degree_type", ""))
    _default_body = 30000 if "硕士" in _degree_type else 50000
    body_target = int(args.body_target if args.body_target is not None else targets.get("body_target_chars", _default_body))
    review_in_scope = bool(args.review_in_scope or targets.get("review_in_scope", False))
    review_target = int(args.review_target if args.review_target is not None else targets.get("review_target_chars", 0))
    references_min_count = int(
        args.references_min if args.references_min is not None else targets.get("references_min_count", 80)
    )
    min_chapters = int(args.min_chapters if args.min_chapters is not None else targets.get("min_chapters", 5))
    per_chapter_ref_floor = targets.get("per_chapter_ref_floor")
    enforce_full_structure = bool(args.enforce_full_structure or ("完整博士论文" in os.path.basename(docx_path)))

    # 生成报告
    report = generate_quality_report(
        docx_path,
        verbose=verbose,
        body_target_chars=body_target,
        review_target_chars=review_target,
        review_in_scope=review_in_scope,
        references_min_count=references_min_count,
        min_chapters=min_chapters,
        enforce_full_structure=enforce_full_structure,
        md_path=args.md_path,
        per_chapter_ref_floor=per_chapter_ref_floor,
        md_checks=args.md_checks,
    )
    
    if not report.get('success'):
        if output_format == 'json':
            print(json.dumps(report, ensure_ascii=False, indent=2))
        else:
            print(f"❌ 检查失败：{report.get('message') or report.get('error')}")
        sys.exit(1)
    
    # 输出报告
    if output_format == 'json':
        print(json.dumps(report, ensure_ascii=False, indent=2))
    else:
        print(format_report_text(report))
    
    # 保存报告
    report_path = docx_path.replace('.docx', '_质量报告.json')
    with open(report_path, 'w', encoding='utf-8') as f:
        json.dump(report, f, ensure_ascii=False, indent=2)
    if verbose:
        print(f"\n💾 报告已保存：{report_path}")
    
    # 返回退出码
    # 硬阻断项：上下标裸写（subsup_bare）+ 中文句内半角标点（halfwidth_punct_in_cn）+ 英文拼写（english_misspelling）
    # + 去AI必禁三项：破折号（decorative_em_dash）/ scare quotes / 解释性冒号（均禁止使用）。
    # 命中即非零退出，无论总分高低；其余 F1 中文错别字仍为 warning（含登陆等同形异义，硬拦会误伤），不在此处阻断。
    subsup_bare_count = report.get('issue_summary', {}).get('subsup_bare', 0)
    halfwidth_in_cn_count = report.get('issue_summary', {}).get('halfwidth_punct_in_cn', 0)
    english_misspelling_count = report.get('issue_summary', {}).get('english_misspelling', 0)
    em_dash_count = report.get('issue_summary', {}).get('decorative_em_dash', 0)
    scare_quote_count = report.get('issue_summary', {}).get('scare_quotes', 0)
    explanatory_colon_count = report.get('issue_summary', {}).get('explanatory_colon', 0)
    if (report['overall_score'] >= 80 and subsup_bare_count == 0 and halfwidth_in_cn_count == 0
            and english_misspelling_count == 0 and em_dash_count == 0
            and scare_quote_count == 0 and explanatory_colon_count == 0):
        sys.exit(0)
    else:
        sys.exit(1)


if __name__ == '__main__':
    main()
